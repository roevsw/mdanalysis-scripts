"""Reusable methods-text generation and Word export utilities.

This module is designed to be extensible: add new builder functions and
register them in `get_default_methods_writer()`.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from datetime import datetime
from html import escape
from pathlib import Path
from typing import Any, Callable, Dict, Optional, Tuple

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from IPython.display import HTML, Markdown, display

Builder = Callable[[Dict[str, Any]], str]


def _fmt(value: Any) -> str:
    """Return a compact numeric string when possible."""
    try:
        return f"{float(value):g}"
    except Exception:
        return str(value)


def _sanitize_xml_text(value: str) -> str:
    """Remove XML-incompatible control characters before writing .docx."""
    return "".join(ch for ch in value if ch in "\t\n\r" or ord(ch) >= 32)


def build_water_spatial_methods_text(params: Dict[str, Any]) -> str:
    """Build methods text for calculate_water_spatial_distribution_xy as prose paragraphs.

    Follows Schimel LD (Lead → Development) structure: flowing prose paragraphs,
    no numbered subsection headers, equations referenced inline, past tense throughout.
    """
    z_centers = params.get("z_slice_centers", [])
    z_slice_width = params.get("z_slice_width", "N/A")
    xy_grid_size = params.get("xy_grid_size", "N/A")
    water_vdw_radius = params.get("water_vdw_radius", "N/A")
    hydrogen_vdw_radius = params.get("hydrogen_vdw_radius", "N/A")
    step = params.get("step", "N/A")

    z_centers_str = ", ".join(_fmt(v) for v in z_centers) if z_centers else "N/A"

    lines = [
        "## Methods: Three-Atom Electron-Weighted Disk Projection",
        "",
        # --- Paragraph 1: overview and slab assignment ---
        rf"The projected areal water density $\rho_{{\mathrm{{proj}}}}(x_i, y_j;\, z_k)$ was computed "
        rf"using a three-atom electron-weighted disk-projection algorithm. "
        rf"Water molecules were assigned to discrete $z$-slabs of thickness "
        rf"$\Delta z = {_fmt(z_slice_width)}$ Å centred at $z_k \in \{{{z_centers_str}\}}$ Å, "
        r"based on the $z$-coordinate of the oxygen atom. A molecule $m$ was included in slab $k$ "
        r"when its oxygen coordinate satisfied",
        "",
        "$$",
        r"\left| z_O^{(m)} - z_k \right| \le \frac{\Delta z}{2}",
        "$$",
        "",
        rf"For each slab, a uniform 2D grid with spacing $\Delta_{{xy}} = {_fmt(xy_grid_size)}$ Å "
        r"was constructed over the $xy$ extent of the simulation box, "
        r"yielding a set of equally spaced grid points $\{(x_i, y_j)\}$.",
        "",
        # --- Paragraph 2: atom disks and electron weights ---
        rf"Each atom $\alpha \in \{{\mathrm{{O}}, \mathrm{{H}}_1, \mathrm{{H}}_2\}}$ of a water molecule "
        rf"was represented as a disk of radius $r_\alpha$ projected onto the $xy$ plane, "
        rf"where $r_O = {_fmt(water_vdw_radius)}$ Å and $r_H = {_fmt(hydrogen_vdw_radius)}$ Å. "
        r"A grid cell $(x_i, y_j)$ was marked as occupied by atom $\alpha$ when the minimum-image "
        r"distance between the cell centre and the atom's $xy$ projection satisfied "
        r"$|\mathbf{r}_{ij} - \mathbf{r}_\alpha^{xy}|_{\mathrm{MIC}} \le r_\alpha$. "
        r"Electron weights ($w_O = 8$, $w_H = 1$, $w_{\mathrm{tot}} = 10$) were assigned to reflect "
        r"the relative scattering contribution of each atom type. "
        r"The weighted areal contribution from atom $\alpha$ was",
        "",
        "$$",
        r"C_\alpha(x_i, y_j) = \frac{w_\alpha}{\pi r_\alpha^2}\,\chi_\alpha(x_i, y_j)",
        "$$",
        "",
        r"where $\chi_\alpha(x_i, y_j) \in \{0, 1\}$ is the disk-occupancy indicator function.",
        "",
        # --- Paragraph 3: per-molecule contribution and frame averaging ---
        r"The per-molecule contribution to each grid cell was obtained by summing the "
        r"electron-weighted atom contributions and normalising by $w_{\mathrm{tot}}$:",
        "",
        "$$",
        r"c_m(x_i, y_j) = \frac{1}{10}\left["
        r"\frac{8\,\chi_O(x_i, y_j)}{\pi r_O^2}"
        r"+ \frac{\chi_{H_1}(x_i, y_j)}{\pi r_H^2}"
        r"+ \frac{\chi_{H_2}(x_i, y_j)}{\pi r_H^2}\right]",
        "$$",
        "",
        rf"The projected areal density was then computed as the average over all $N_f$ trajectory frames, "
        rf"sampled every $\mathrm{{step}} = {_fmt(step)}$ frames, by accumulating per-molecule contributions "
        r"from all water molecules in slab $k$ across each frame:",
        "",
        "$$",
        r"\rho_{\mathrm{proj}}(x_i, y_j;\, z_k) = "
        r"\frac{1}{N_f}\sum_{f=1}^{N_f}\sum_{m \in \mathrm{slab}_k,\,f} c_m(x_i, y_j)",
        "$$",
        "",
        r"The resulting $\rho_{\mathrm{proj}}$ is reported in units of molecules/Å$^2$ "
        r"(projected 2D areal density).",
    ]

    return "\n".join(lines)


def build_water_spatial_methods_manuscript_blocks(params: Dict[str, Any]):
    """Build manuscript-friendly prose blocks for Word output (no raw LaTeX).

    Follows Schimel LD structure: three flowing paragraphs with equations
    referenced inline by number. No numbered subsection headers. Past tense.
    """
    z_centers = params.get("z_slice_centers", [])
    z_slice_width = params.get("z_slice_width", "N/A")
    xy_grid_size = params.get("xy_grid_size", "N/A")
    water_vdw_radius = params.get("water_vdw_radius", "N/A")
    hydrogen_vdw_radius = params.get("hydrogen_vdw_radius", "N/A")
    step = params.get("step", "N/A")

    z_centers_str = ", ".join(_fmt(v) for v in z_centers) if z_centers else "N/A"

    return [
        # --- Paragraph 1: overview and slab assignment ---
        (
            "p",
            f"The projected areal water density \u03c1_proj(x\u1d62, y\u2c7c;\u2003z_k) was computed "
            f"using a three-atom electron-weighted disk-projection algorithm. "
            f"Water molecules were assigned to discrete z-slabs of thickness "
            f"\u0394z\u2009=\u2009{_fmt(z_slice_width)} \u00c5 centred at "
            f"z_k \u2208 {{{z_centers_str}}} \u00c5, based on the z-coordinate of the oxygen atom. "
            f"A molecule m was included in slab k when its oxygen coordinate satisfied "
            f"|z_O(m) \u2212 z_k| \u2264 \u0394z/2 (Eq.\u20091). "
            f"For each slab, a uniform 2D grid with spacing "
            f"\u0394_xy\u2009=\u2009{_fmt(xy_grid_size)} \u00c5 was constructed over the xy extent "
            f"of the simulation box, yielding a set of equally spaced grid points {{(x\u1d62, y\u2c7c)}}.",
        ),
        ("eq", "(Eq.\u20091)\u2003|z_O(m) \u2212 z_k| \u2264 \u0394z/2"),
        # --- Paragraph 2: atom disks and electron weighting ---
        (
            "p",
            f"Each atom \u03b1 \u2208 {{O, H\u2081, H\u2082}} of a water molecule was represented as a disk "
            f"of radius r_\u03b1 projected onto the xy plane, where r_O\u2009=\u2009{_fmt(water_vdw_radius)} \u00c5 "
            f"and r_H\u2009=\u2009{_fmt(hydrogen_vdw_radius)} \u00c5. "
            f"A grid cell (x\u1d62, y\u2c7c) was marked as occupied by atom \u03b1 when the "
            f"minimum-image distance between the cell centre and the atom\u2019s xy projection "
            f"was at most r_\u03b1 (Eq.\u20092). Electron weights (w_O\u2009=\u20098, w_H\u2009=\u20091, "
            f"w_tot\u2009=\u200910) were assigned to reflect the relative scattering contribution of each "
            f"atom type. The weighted areal contribution from atom \u03b1 to a grid cell was "
            f"computed as shown in Eq.\u20093.",
        ),
        ("eq", "(Eq.\u20092)\u2003\u03c7_\u03b1(x\u1d62, y\u2c7c) = 1 if |r_ij \u2212 r_\u03b1^xy|_MIC \u2264 r_\u03b1, else 0"),
        ("eq", "(Eq.\u20093)\u2003C_\u03b1(x\u1d62, y\u2c7c) = [w_\u03b1\u2009/\u2009(\u03c0 r_\u03b1\u00b2)] \u00b7 \u03c7_\u03b1(x\u1d62, y\u2c7c)"),
        # --- Paragraph 3: per-molecule contribution and frame averaging ---
        (
            "p",
            f"The per-molecule contribution to each grid cell was obtained by summing the "
            f"electron-weighted atom contributions and normalising by the total electron count "
            f"of a water molecule (w_tot\u2009=\u200910), as given by Eq.\u20094. The projected areal "
            f"water density \u03c1_proj was then computed as the average over all N_f trajectory "
            f"frames, sampled every {_fmt(step)} frames, by accumulating per-molecule "
            f"contributions from all water molecules in slab k across each frame (Eq.\u20095). "
            f"The resulting \u03c1_proj is reported in units of molecules/\u00c5\u00b2 "
            f"(projected 2D areal density).",
        ),
        ("eq", "(Eq.\u20094)\u2003c_m(x\u1d62, y\u2c7c) = (1/10) \u00b7 [8\u03c7_O/(\u03c0r_O\u00b2) + \u03c7_H\u2081/(\u03c0r_H\u00b2) + \u03c7_H\u2082/(\u03c0r_H\u00b2)]"),
        ("eq", "(Eq.\u20095)\u2003\u03c1_proj(x\u1d62, y\u2c7c;\u2003z_k) = (1/N_f) \u00b7 \u2211_f \u2211_{m\u2208slab_k} c_m(x\u1d62, y\u2c7c)"),
    ]


@dataclass
class MethodsWriter:
    """Generic writer for notebook display + Word export from registered builders."""

    builders: Dict[str, Builder] = field(default_factory=dict)

    def register_builder(self, name: str, builder: Builder) -> None:
        self.builders[name] = builder

    def generate(
        self,
        method_name: str,
        params: Dict[str, Any],
        output_dir: Optional[Path] = None,
        show_in_notebook: bool = True,
        save_docx: bool = True,
        filename_prefix: Optional[str] = None,
        doc_title: Optional[str] = None,
    ) -> Tuple[str, Optional[Path]]:
        if method_name not in self.builders:
            available = ", ".join(sorted(self.builders)) or "<none>"
            raise ValueError(f"Unknown method '{method_name}'. Available: {available}")

        text = self.builders[method_name](params)
        text = _sanitize_xml_text(text)

        if show_in_notebook:
            self._show_text(text)

        saved_path = None
        if save_docx:
            out_dir = Path(output_dir) if output_dir else Path.cwd()
            out_dir.mkdir(parents=True, exist_ok=True)

            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            prefix = filename_prefix or method_name
            saved_path = out_dir / f"{prefix}_{ts}.docx"

            title = doc_title or method_name.replace("_", " ").title()
            self._save_docx(text, saved_path, title)
            print(f"Word file saved -> {saved_path}")

        return text, saved_path

    @staticmethod
    def _show_text(text: str) -> None:
        # Render as Markdown so LaTeX equations display as real math in notebooks.
        display(Markdown(text))

    @staticmethod
    def _save_docx(text: str, out_path: Path, title: str) -> None:
        doc = Document()

        # Force heading colors to black to avoid Word theme blue defaults.
        for heading_style_name in ("Heading 1", "Heading 2"):
            if heading_style_name in doc.styles:
                doc.styles[heading_style_name].font.name = "Times New Roman"
                doc.styles[heading_style_name].font.color.rgb = RGBColor(0, 0, 0)

        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            run.font.name = "Times New Roman"
            run.font.color.rgb = RGBColor(0, 0, 0)

        for block in text.split("\n\n"):
            line = _sanitize_xml_text(block.strip())
            if not line:
                continue
            if line.startswith("## "):
                continue
            if line.startswith("### "):
                h2 = doc.add_heading(line[4:].strip(), level=2)
                for run in h2.runs:
                    run.font.name = "Times New Roman"
                    run.font.color.rgb = RGBColor(0, 0, 0)
                continue

            p = doc.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(8)
            for run in p.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)

        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)

        doc.save(str(out_path))


def get_default_methods_writer() -> MethodsWriter:
    """Return a writer preconfigured with built-in builders."""
    writer = MethodsWriter()
    writer.register_builder("water_spatial_projection_methods", build_water_spatial_methods_text)
    writer.register_builder("water_dipole_z_methods", build_water_dipole_z_methods_text)
    writer.register_builder("water_dipole_xy_methods", build_water_dipole_xy_methods_text)
    writer.register_builder("ion_spatial_xy_methods", build_ion_spatial_xy_methods_text)
    writer.register_builder(
        "electrostatic_potential_methods", build_electrostatic_potential_methods_text
    )
    writer.register_builder(
        "coordination_shells_methods", build_coordination_shells_methods_text
    )
    writer.register_builder(
        "rdf_vs_z_methods", build_rdf_vs_z_methods_text
    )
    return writer


def write_water_spatial_methods(
    results_obj: Any = None,
    params: Optional[Dict[str, Any]] = None,
    output_dir: Optional[Path] = None,
    show_in_notebook: bool = True,
    save_docx: bool = True,
):
    """Compatibility wrapper for notebook use.

    results_obj is accepted for API stability, even if unused in this text builder.
    """
    del results_obj

    if params is None:
        params = globals().get("water_spatial_kwargs", {})

    if not params:
        raise ValueError("No water_spatial_kwargs found. Run the calculation cell first.")

    markdown_text = build_water_spatial_methods_text(params)
    if show_in_notebook:
        blocks = build_water_spatial_methods_manuscript_blocks(params)
        _show_methods_as_white_page(
            title="Methods: Three-Atom Electron-Weighted Disk Projection",
            blocks=blocks,
        )

    saved_path = None
    if save_docx:
        out_dir = Path(output_dir) if output_dir else Path.cwd()
        out_dir.mkdir(parents=True, exist_ok=True)

        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        saved_path = out_dir / f"water_spatial_projection_methods_{ts}.docx"

        doc = Document()

        # Force heading colors to black to avoid Word theme blue defaults.
        for heading_style_name in ("Heading 1", "Heading 2"):
            if heading_style_name in doc.styles:
                doc.styles[heading_style_name].font.name = "Times New Roman"
                doc.styles[heading_style_name].font.color.rgb = RGBColor(0, 0, 0)

        heading = doc.add_heading("Methods: Three-Atom Electron-Weighted Disk Projection", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            run.font.name = "Times New Roman"
            run.font.color.rgb = RGBColor(0, 0, 0)

        for kind, text in build_water_spatial_methods_manuscript_blocks(params):
            line = _sanitize_xml_text(text)
            if not line:
                continue

            if kind == "h2":
                h2 = doc.add_heading(line, level=2)
                for run in h2.runs:
                    run.font.name = "Times New Roman"
                    run.font.color.rgb = RGBColor(0, 0, 0)
                continue

            p = doc.add_paragraph(line)
            if kind == "eq":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.space_after = Pt(8)
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)

        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)

        doc.save(str(saved_path))
        print(f"Word file saved -> {saved_path}")

    return markdown_text, saved_path


def _show_methods_as_white_page(title: str, blocks) -> None:
        """Render manuscript text as a white page in notebook output (theme-independent)."""
        html_parts = [
                """
<style>
.methods-page {
    max-width: 900px;
    margin: 16px 0;
    padding: 34px 42px;
    border: 1px solid #d8d8d8;
    border-radius: 4px;
    background: #ffffff !important;
    color: #111111 !important;
    box-shadow: 0 2px 10px rgba(0,0,0,0.08);
    font-family: "Times New Roman", Georgia, serif;
    line-height: 1.6;
}
.methods-page h1 {
    margin: 0 0 16px 0;
    text-align: left;
    font-size: 24px;
    font-weight: 700;
    color: #111111;
}
.methods-page h2 {
    margin: 18px 0 8px 0;
    font-size: 18px;
    font-weight: 700;
    color: #111111;
}
.methods-page p {
    margin: 8px 0;
    font-size: 15px;
    color: #111111;
    text-align: justify;
}
.methods-page .eq {
    margin: 10px 0;
    text-align: center;
    font-size: 16px;
    color: #111111;
    font-family: "Cambria Math", "STIX Two Math", "Times New Roman", serif;
}
</style>
"""
        ]

        html_parts.append("<div class='methods-page'>")
        html_parts.append(f"<h1>{escape(title)}</h1>")

        for kind, text in blocks:
                safe = escape(str(text))
                if kind == "h2":
                        html_parts.append(f"<h2>{safe}</h2>")
                elif kind == "eq":
                        html_parts.append(f"<div class='eq'>{safe}</div>")
                else:
                        html_parts.append(f"<p>{safe}</p>")

        html_parts.append("</div>")
        display(HTML("".join(html_parts)))


# ──────────────────────────────────────────────────────────────────────────────
# WATER DIPOLE ORIENTATION — Z-PROFILE (2D Heatmap)
# ──────────────────────────────────────────────────────────────────────────────

def build_water_dipole_z_methods_text(params: Dict[str, Any]) -> str:
    """Build Markdown/LaTeX methods text for water dipole orientation vs z (2D heatmap).

    Follows Schimel LD structure: two flowing prose paragraphs, equations
    referenced inline, past tense throughout.
    """
    bin_size = params.get("bin_size", 2.0)
    angle_bin_size = params.get("angle_bin_size", 5)

    lines = [
        "## Methods: Water Dipole Orientation Along the Channel Normal",
        "",
        # --- Paragraph 1: dipole vector definition and angle ---
        r"The orientation of water molecule dipoles as a function of position along the "
        r"channel normal ($z$-axis) was characterised by computing, for every water "
        r"molecule in every trajectory frame, the molecular dipole vector $\mathbf{d}$ "
        r"defined as the displacement from the mean hydrogen position to the oxygen atom:",
        "",
        "$$",
        r"\mathbf{d} = \mathbf{r}_O - \tfrac{1}{2}\left(\mathbf{r}_{H_1} + \mathbf{r}_{H_2}\right)",
        "$$",
        "",
        r"This convention aligns $\mathbf{d}$ with the direction of the permanent molecular "
        r"dipole moment (from the partial-positive hydrogen end toward the partial-negative "
        r"oxygen). The angle $\theta$ between the unit dipole vector $\hat{\mathbf{d}}$ and "
        r"the positive $z$-axis was then obtained as",
        "",
        "$$",
        r"\theta = \arccos\!\left(\hat{\mathbf{d}} \cdot \hat{\mathbf{z}}\right), "
        r"\quad \hat{\mathbf{d}} = \frac{\mathbf{d}}{|\mathbf{d}|}",
        "$$",
        "",
        r"yielding values in the range $0^\circ$–$180^\circ$. A value of $\theta = 90^\circ$ "
        r"corresponds to a randomly oriented dipole; $\theta < 90^\circ$ indicates preferential "
        r"alignment of the oxygen toward the positive-$z$ clay surface; $\theta > 90^\circ$ "
        r"signals preferential anti-alignment.",
        "",
        # --- Paragraph 2: 2D histogram accumulation ---
        rf"The joint distribution $P(z,\,\theta)$ was accumulated as a two-dimensional "
        rf"histogram over all water molecules and all trajectory frames, with $z$-bins of "
        rf"width $\Delta z_\mathrm{{bin}} = {_fmt(bin_size)}$ Å and angular bins of "
        rf"$\Delta\theta = {_fmt(angle_bin_size)}^\circ$. Each bin records the total count "
        r"of molecule–frame observations falling within the corresponding "
        r"$(z,\,\theta)$ interval, reflecting sampling frequency across the trajectory "
        r"rather than a normalised probability density. The $z$-coordinate of the oxygen "
        r"atom was used to assign each molecule to its $z$-bin, after applying a "
        r"box-centring offset so that $z = 0$ corresponds to the geometric centre of the "
        r"simulation cell.",
        "",
        # --- Paragraph 3: top vs bottom surface interpretation ---
        r"Because the histogram spans the full $z$-range of the simulation cell, it "
        r"captures the dipole distribution near both clay\u2013water interfaces "
        r"simultaneously. With $z = 0$ at the channel midpoint, the upper clay surface "
        r"occupies $z > 0$ and the lower clay surface occupies $z < 0$. The angle "
        r"$\theta$ is referenced to the same global $+z$ direction throughout, so its "
        r"physical meaning is surface-dependent: near the upper surface ($z > 0$), "
        r"$\theta < 90^\circ$ indicates that the oxygen end of the dipole is directed "
        r"toward the clay above; near the lower surface ($z < 0$), the clay lies in the "
        r"$-z$ direction, so the same oxygen-toward-clay orientation corresponds to "
        r"$\theta > 90^\circ$. Consequently, if both surfaces are structurally "
        r"equivalent, the angular distributions at symmetric positions $+z_k$ and "
        r"$-z_k$ are mirror-complementary about $\theta = 90^\circ$.",
    ]
    return "\n".join(lines)


def build_water_dipole_z_methods_manuscript_blocks(params: Dict[str, Any]):
    """Build Word-friendly Unicode prose blocks for dipole-vs-z methods.

    Two LD-structure paragraphs, equations referenced inline, no sub-headers.
    """
    bin_size = params.get("bin_size", 2.0)
    angle_bin_size = params.get("angle_bin_size", 5)

    return [
        (
            "p",
            "The orientation of water molecule dipoles as a function of position along "
            "the channel normal (z-axis) was characterised by computing, for every water "
            "molecule in every trajectory frame, the molecular dipole vector d defined as "
            "the displacement from the mean hydrogen position to the oxygen atom (Eq.\u20091). "
            "This convention aligns d with the direction of the permanent molecular dipole "
            "moment, pointing from the partial-positive hydrogen end toward the "
            "partial-negative oxygen. The angle \u03b8 between the unit dipole vector d\u0302 "
            "and the positive z-axis was obtained as \u03b8 = arccos(d\u0302 \u00b7 \u1e91) "
            "(Eq.\u20092), yielding values in the range 0\u00b0\u2013180\u00b0. "
            "A value of \u03b8 = 90\u00b0 corresponds to a randomly oriented dipole; "
            "\u03b8 < 90\u00b0 indicates preferential alignment of the oxygen toward the "
            "positive-z clay surface; \u03b8 > 90\u00b0 signals preferential anti-alignment.",
        ),
        ("eq", "(Eq.\u20091)\u2003d = r_O \u2212 \u00bd(r_H\u2081 + r_H\u2082)"),
        ("eq", "(Eq.\u20092)\u2003\u03b8 = arccos(d\u0302 \u00b7 \u1e91),\u2003d\u0302 = d / |d|"),
        (
            "p",
            f"The joint distribution P(z, \u03b8) was accumulated as a two-dimensional "
            f"histogram over all water molecules and all trajectory frames, with z-bins of "
            f"width \u0394z\u209c = {_fmt(bin_size)} \u00c5 and angular bins of "
            f"\u0394\u03b8 = {_fmt(angle_bin_size)}\u00b0. Each bin records the total count "
            "of molecule\u2013frame observations falling within the corresponding "
            "(z, \u03b8) interval, reflecting sampling frequency across the trajectory "
            "rather than a normalised probability density. The z-coordinate of the oxygen "
            "atom was used to assign each molecule to its z-bin, after applying a "
            "box-centring offset so that z = 0 corresponds to the geometric centre of the "
            "simulation cell.",
        ),
        (
            "p",
            "Because the histogram spans the full z-range of the simulation cell, it "
            "captures the dipole distribution near both clay\u2013water interfaces "
            "simultaneously. With z\u00a0=\u00a00 at the channel midpoint, the upper clay "
            "surface occupies z\u00a0>\u00a00 and the lower clay surface occupies z\u00a0<\u00a00. "
            "The angle \u03b8 is referenced to the same global +z direction throughout, so "
            "its physical meaning is surface-dependent: near the upper surface (z\u00a0>\u00a00), "
            "\u03b8\u00a0<\u00a090\u00b0 indicates that the oxygen end of the dipole is directed toward "
            "the clay above; near the lower surface (z\u00a0<\u00a00), the clay lies in the \u2212z "
            "direction, so the same oxygen-toward-clay orientation corresponds to "
            "\u03b8\u00a0>\u00a090\u00b0. Consequently, if both surfaces are structurally equivalent, the "
            "angular distributions at symmetric positions +z\u2096 and \u2212z\u2096 are "
            "mirror-complementary about \u03b8\u00a0=\u00a090\u00b0.",
        ),
    ]


# ──────────────────────────────────────────────────────────────────────────────
# WATER DIPOLE ORIENTATION — XY PLANAR MAPS
# ──────────────────────────────────────────────────────────────────────────────

def _dipole_xy_spatial_paragraph_text(params: Dict[str, Any]) -> tuple:
    """Return (markdown_paragraph, [unicode_blocks]) for the spatial-assignment step."""
    xy_projection = params.get("xy_projection", False)
    smoothing_method = params.get("smoothing_method", "none")
    projection_method = params.get("projection_method", "elliptical")
    water_vdw_radius = params.get("water_vdw_radius", 1.52)
    hydrogen_vdw_radius = params.get("hydrogen_vdw_radius", 1.2)
    projection_radius = params.get("projection_radius", 2.8)
    ellipse_aspect_ratio = params.get("ellipse_aspect_ratio", 2.0)

    if xy_projection:
        if projection_method == "three_atom":
            md = (
                r"The molecular footprint on the $xy$ grid was assigned using a three-atom "
                r"disk-projection scheme. Each of the three atoms — O, H$_1$, H$_2$ — "
                r"was projected as a filled disk of its respective van der Waals radius "
                rf"($r_O = {_fmt(water_vdw_radius)}$ Å, $r_H = {_fmt(hydrogen_vdw_radius)}$ Å). "
                r"Grid cells within the disk of atom $\alpha$ received a contribution "
                r"proportional to the atomic electron count ($w_O = 8$, $w_H = 1$), "
                r"normalised by the total electron count per molecule ($w_\mathrm{tot} = 10$). "
                r"The orientation value was then distributed across all occupied cells in "
                r"proportion to these normalised weights, ensuring that the molecular footprint "
                r"reflects the electron-density distribution of the water molecule."
            )
            ub = [(
                "p",
                f"The molecular footprint on the xy grid was assigned using a three-atom "
                f"disk-projection scheme. Each atom (O, H\u2081, H\u2082) was projected as a "
                f"filled disk of its van der Waals radius "
                f"(r_O = {_fmt(water_vdw_radius)} \u00c5, r_H = {_fmt(hydrogen_vdw_radius)} \u00c5). "
                "Grid cells within each disk received a contribution weighted by the atomic "
                "electron count (w_O = 8, w_H = 1), normalised by the total electron count "
                "per molecule (w_tot = 10). The orientation value was distributed across all "
                "occupied cells in proportion to these normalised weights, so that the "
                "molecular footprint reflects the electron-density profile of the water molecule.",
            )]
        elif projection_method == "elliptical":
            b = _fmt(projection_radius)
            a = _fmt(float(projection_radius) * float(ellipse_aspect_ratio))
            md = (
                r"The molecular footprint was projected onto the $xy$ plane as a "
                r"dipole-oriented ellipse. The major axis of the ellipse was aligned with "
                r"the $xy$ projection of the dipole vector $\hat{\mathbf{d}}$, with "
                rf"semi-major length $a = {a}$ Å and semi-minor length $b = {b}$ Å "
                rf"(aspect ratio $a/b = {_fmt(ellipse_aspect_ratio)}$). "
                r"All grid cells falling within the ellipse received a uniform contribution "
                r"from that molecule, normalised by the number of occupied cells. "
                r"For dipoles with a negligible $xy$ component ($|\hat{\mathbf{d}}_{xy}| < 0.01$), "
                r"a circular footprint of radius $b$ was substituted as a fallback."
            )
            ub = [(
                "p",
                f"The molecular footprint was projected onto the xy plane as a "
                f"dipole-oriented ellipse. The major axis was aligned with the xy projection "
                f"of the dipole vector, with semi-major length a = {a} \u00c5 "
                f"and semi-minor length b = {b} \u00c5 (aspect ratio {_fmt(ellipse_aspect_ratio)}). "
                "All grid cells within the ellipse received a uniform contribution, normalised "
                "by the cell count. For dipoles with a negligible xy component, a circular "
                f"footprint of radius b = {b} \u00c5 was used as a fallback.",
            )]
        else:  # circular
            md = (
                rf"Each water molecule was projected onto the $xy$ grid as a uniform disk of "
                rf"radius $r_\mathrm{{proj}} = {_fmt(projection_radius)}$ Å centred on the "
                r"oxygen position. All grid cells within this disk received an equal "
                r"contribution from that molecule, normalised by the number of occupied cells."
            )
            ub = [(
                "p",
                f"Each water molecule was projected onto the xy grid as a uniform disk of "
                f"projected radius r_proj = {_fmt(projection_radius)} \u00c5 centred on the oxygen "
                "position. All grid cells within this disk received an equal contribution, "
                "normalised by the cell count.",
            )]
    else:
        if smoothing_method == "none":
            md = (
                r"Each water molecule contributed its orientation value to a single grid cell "
                r"determined by digitising the periodic oxygen $xy$ position into the "
                r"nearest bin of the grid. No spatial smoothing was applied; each grid cell "
                r"therefore accumulated the raw count and summed orientation value of all "
                r"molecules assigned to it across all sampled frames."
            )
            ub = [(
                "p",
                "Each water molecule contributed its orientation value to a single grid cell "
                "determined by digitising the periodic oxygen xy position into the nearest bin. "
                "No spatial smoothing was applied; each grid cell accumulated the raw count "
                "and orientation sum of all molecules assigned to it across all sampled frames.",
            )]
        elif smoothing_method == "vdw":
            md = (
                rf"Each water molecule was spread over the $xy$ grid using a Gaussian "
                rf"weight profile centred on the oxygen position, with width parameter "
                rf"$\sigma = r_O = {_fmt(water_vdw_radius)}$ Å equal to the oxygen van der Waals "
                r"radius. Contributions were truncated at a cutoff of $3\sigma$. "
                r"The orientation value was accumulated as the product of this weight "
                r"and the scalar orientation metric at each grid cell."
            )
            ub = [(
                "p",
                f"Each water molecule was spread over the xy grid using a Gaussian weight "
                f"profile centred on the oxygen position, with \u03c3 = r_O = "
                f"{_fmt(water_vdw_radius)} \u00c5 (oxygen van der Waals radius). Contributions were "
                "truncated at 3\u03c3. The orientation value was accumulated as the product of "
                "the Gaussian weight and the scalar orientation metric at each grid cell.",
            )]
        elif smoothing_method == "molecular_vdw":
            oh_bond = 0.96
            r_mol = float(water_vdw_radius) + oh_bond + float(hydrogen_vdw_radius)
            md = (
                rf"Spatial smoothing used the effective molecular van der Waals radius "
                rf"$r_\mathrm{{mol}} = r_O + d_{{OH}} + r_H = "
                rf"{_fmt(water_vdw_radius)} + {_fmt(oh_bond)} + {_fmt(hydrogen_vdw_radius)} "
                rf"= {r_mol:.2f}$ Å, treating the water molecule as a single composite sphere. "
                r"A Gaussian weight centred on the oxygen position with $\sigma = r_\mathrm{mol}$ "
                r"and a cutoff at $3\sigma$ was used to distribute the orientation value "
                r"across the grid."
            )
            ub = [(
                "p",
                f"Spatial smoothing used the effective molecular van der Waals radius "
                f"r_mol = r_O + d_OH + r_H = {_fmt(water_vdw_radius)} + {_fmt(oh_bond)} + "
                f"{_fmt(hydrogen_vdw_radius)} = {r_mol:.2f} \u00c5, treating the water molecule as "
                "a single composite sphere. A Gaussian weight centred on the oxygen position "
                "with \u03c3 = r_mol and a cutoff at 3\u03c3 was applied to distribute the "
                "orientation value across the grid.",
            )]
        elif smoothing_method == "three_point":
            md = (
                r"Spatial smoothing was performed at all three atomic positions (O, H$_1$, H$_2$) "
                r"simultaneously. A Gaussian weight with atom-specific width "
                rf"($\sigma_O = {_fmt(water_vdw_radius)}$ Å, "
                rf"$\sigma_H = {_fmt(hydrogen_vdw_radius)}$ Å) "
                r"and a $3\sigma$ hard cutoff was centred at each atom. Atomic contributions "
                r"were weighted by physical importance factors ($w_O = 2.0$, $w_H = 0.5$) "
                r"and summed; the result was normalised by the total weight ($w_\mathrm{tot} = 3.0$) "
                r"before being added to the orientation and count grids."
            )
            ub = [(
                "p",
                f"Spatial smoothing was performed at all three atomic positions (O, H\u2081, H\u2082) "
                f"simultaneously. A Gaussian weight with atom-specific width "
                f"(\u03c3_O = {_fmt(water_vdw_radius)} \u00c5, "
                f"\u03c3_H = {_fmt(hydrogen_vdw_radius)} \u00c5) "
                "and a 3\u03c3 hard cutoff was centred at each atom. Contributions were "
                "weighted by physical importance factors (w_O = 2.0, w_H = 0.5), summed, "
                "and normalised by the total weight (w_tot = 3.0) before accumulation.",
            )]
        else:  # elliptical
            md = (
                r"Smoothing used a dipole-oriented elliptical Gaussian centred on the oxygen, "
                rf"with major semi-axis $a = r_O \times {_fmt(ellipse_aspect_ratio)} $ Å along the "
                r"$xy$ projection of the dipole and minor semi-axis $b = r_O$ perpendicular to it. "
                r"A $3\sigma$ elliptical cutoff (in units of $a$ and $b$ respectively) was applied. "
                r"For dipoles with negligible $xy$ component, a circular Gaussian of width "
                r"$\sigma = r_O$ was substituted."
            )
            ub = [(
                "p",
                f"Smoothing used a dipole-oriented elliptical Gaussian centred on the oxygen, "
                f"with major semi-axis a = r_O \u00d7 {_fmt(ellipse_aspect_ratio)} \u00c5 along the "
                "xy projection of the dipole and minor semi-axis b = r_O perpendicular to it. "
                "A 3\u03c3 elliptical cutoff was applied. For dipoles with a negligible xy "
                "component, a circular Gaussian of width \u03c3 = r_O was substituted.",
            )]

    return md, ub


def _dipole_xy_metric_sentence(orientation_metric: str) -> tuple:
    """Return (markdown_sentence, unicode_sentence) for the orientation metric."""
    if orientation_metric == "angle_to_z":
        md = (
            r"The orientation metric mapped to each grid cell was the dipole angle "
            r"$\theta = \arccos(\hat{\mathbf{d}} \cdot \hat{\mathbf{z}})$ in degrees "
            r"(Eq.\,3), ranging from $0^\circ$ (oxygen pointing along +z) to "
            r"$180^\circ$ (oxygen pointing along \textminus z), with $90^\circ$ "
            r"representing a horizontal, randomly oriented dipole."
        )
        ub = (
            "The orientation metric mapped to each grid cell was the dipole angle "
            "\u03b8 = arccos(d\u0302 \u00b7 \u1e91) in degrees (Eq.\u20093), ranging from "
            "0\u00b0 (oxygen pointing along +z) to 180\u00b0 (oxygen pointing along \u2212z), "
            "with 90\u00b0 representing a horizontal, randomly oriented dipole."
        )
        eq_md = r"$$\theta = \arccos\!\left(\hat{\mathbf{d}} \cdot \hat{\mathbf{z}}\right)$$"
        eq_ub = "(Eq.\u20093)\u2003\u03b8 = arccos(d\u0302 \u00b7 \u1e91)"
    elif orientation_metric == "order_parameter":
        md = (
            r"The orientation metric was the second-rank orientational order parameter "
            r"$S = \tfrac{1}{2}(3\cos^2\theta - 1)$ (Eq.\,3), ranging from $-0.5$ "
            r"(dipole perpendicular to $z$) through $0$ (random) to $1.0$ (fully aligned)."
        )
        ub = (
            "The orientation metric was the second-rank orientational order parameter "
            "S = \u00bd(3cos\u00b2\u03b8 \u2212 1) (Eq.\u20093), ranging from \u22120.5 "
            "(dipole perpendicular to z) through 0 (random) to 1.0 (fully aligned)."
        )
        eq_md = r"$$S = \frac{1}{2}\left(3\cos^2\theta - 1\right)$$"
        eq_ub = "(Eq.\u20093)\u2003S = \u00bd(3cos\u00b2\u03b8 \u2212 1)"
    else:  # vector_field
        md = (
            r"The orientation metric was the three Cartesian components of the unit dipole "
            r"vector $\hat{\mathbf{d}} = (d_x,\, d_y,\, d_z)$, each accumulated and "
            r"averaged independently on the grid, yielding a spatially resolved vector field "
            r"of mean dipole orientation (Eq.\,3)."
        )
        ub = (
            "The orientation metric was the three Cartesian components of the unit dipole "
            "vector d\u0302 = (d_x, d_y, d_z), each accumulated and averaged independently, "
            "yielding a spatially resolved vector field of mean dipole orientation (Eq.\u20093)."
        )
        eq_md = r"$$\hat{\mathbf{d}} = \left(d_x,\, d_y,\, d_z\right) / |\mathbf{d}|$$"
        eq_ub = "(Eq.\u20093)\u2003d\u0302 = (d_x, d_y, d_z) / |d|"

    return md, ub, eq_md, eq_ub


def build_water_dipole_xy_methods_text(params: Dict[str, Any]) -> str:
    """Build Markdown/LaTeX methods text for the planar XY dipole orientation maps.

    Three LD-structure prose paragraphs, past tense, equations referenced inline.
    """
    z_centers = params.get("z_slice_centers", [])
    z_slice_width = params.get("z_slice_width", 2.0)
    xy_grid_size = params.get("xy_grid_size", 0.5)
    step = params.get("step", 1)
    orientation_metric = params.get("orientation_metric", "angle_to_z")
    water_vdw_radius = params.get("water_vdw_radius", 1.52)
    hydrogen_vdw_radius = params.get("hydrogen_vdw_radius", 1.2)

    z_centers_str = ", ".join(_fmt(v) for v in z_centers) if z_centers else "auto-detected"
    n_slices = len(z_centers) if z_centers else "N"

    metric_md, _, eq_md, _ = _dipole_xy_metric_sentence(orientation_metric)
    spatial_md, _ = _dipole_xy_spatial_paragraph_text(params)

    lines = [
        "## Methods: Planar Water Dipole Orientation Maps",
        "",
        # --- Paragraph 1: slab geometry and grid setup ---
        rf"The spatial distribution of water dipole orientation in the $xy$ plane was "
        rf"computed at {n_slices} discrete $z$-slices centred at "
        rf"$z_k \in \{{{z_centers_str}\}}$ Å. "
        rf"Within each slab of thickness $\Delta z = {_fmt(z_slice_width)}$ Å, molecules "
        r"whose oxygen $z$-coordinate satisfied",
        "",
        "$$",
        r"\left| z_O - z_k \right| \le \frac{\Delta z}{2}",
        "$$",
        "",
        rf"were selected for analysis (Eq.\,1). A uniform $xy$ grid of spacing "
        rf"$\Delta_{{xy}} = {_fmt(xy_grid_size)}$ Å was constructed over the full "
        r"periodic box, and the dipole vector $\mathbf{d}$ was computed for each "
        r"selected molecule from",
        "",
        "$$",
        r"\mathbf{d} = \mathbf{r}_O - \tfrac{1}{2}\left(\mathbf{r}_{H_1} + "
        r"\mathbf{r}_{H_2}\right)",
        "$$",
        "",
        r"consistent with the direction of the permanent dipole moment (Eq.\,2). "
        + metric_md,
        "",
        eq_md,
        "",
        # --- Paragraph 2: spatial assignment ---
        spatial_md,
        "",
        # --- Paragraph 3: frame averaging ---
        rf"The trajectory was sampled at every {_fmt(step)} frame(s). For each grid cell, "
        r"the final orientation value $\bar{V}(x_i, y_j;\, z_k)$ was obtained as the "
        r"weighted average over all accumulated contributions:",
        "",
        "$$",
        r"\bar{V}(x_i, y_j;\, z_k) = \frac{\displaystyle\sum_{f,m} w_{f,m}(x_i,y_j)\, "
        r"V_{f,m}}{\displaystyle\sum_{f,m} w_{f,m}(x_i,y_j)}",
        "$$",
        "",
        r"where $V_{f,m}$ is the orientation metric of molecule $m$ in frame $f$, "
        r"$w_{f,m}(x_i, y_j)$ is the spatial weight assigned to cell $(x_i, y_j)$, "
        r"and the sums run over all molecules in the slab across all analysed frames. "
        r"Grid cells with no contributing molecules were left at zero and excluded from "
        r"visualisation.",
        "",
        # --- Paragraph 4: top vs bottom surface interpretation ---
        r"Slices were placed near both clay\u2013water interfaces. Where slice centres "
        r"were auto-detected from the clay interface boundaries, slices at the upper "
        r"surface (positive-$z$ side) were positioned at $z_k = z_\mathrm{surf}^+ - \delta$ "
        r"(water immediately below the upper clay), and slices at the lower surface "
        r"(negative-$z$ side) at $z_k = z_\mathrm{surf}^- + \delta$, where "
        r"$\delta \in \{2, 5\}$ \AA{}. The channel midpoint ($z = 0$) was also "
        r"included as a reference slice. Because the dipole angle $\theta$ is always "
        r"measured against the global $+z$ axis, its directional meaning is "
        r"surface-dependent: at slices near the upper clay, $\theta < 90^\circ$ "
        r"indicates preferential oxygen alignment toward that surface; at slices near "
        r"the lower clay, $\theta > 90^\circ$ signals the same orientation relative to "
        r"the surface below. Consequently, maps from the two surfaces display "
        r"complementary colour distributions when the colourmap is centred at $90^\circ$.",
    ]
    return "\n".join(lines)


def build_water_dipole_xy_methods_manuscript_blocks(params: Dict[str, Any]):
    """Build Word-friendly Unicode prose blocks for the XY dipole orientation methods.

    Three LD-structure paragraphs, equations referenced by number, no sub-headers.
    """
    z_centers = params.get("z_slice_centers", [])
    z_slice_width = params.get("z_slice_width", 2.0)
    xy_grid_size = params.get("xy_grid_size", 0.5)
    step = params.get("step", 1)
    orientation_metric = params.get("orientation_metric", "angle_to_z")

    z_centers_str = ", ".join(_fmt(v) for v in z_centers) if z_centers else "auto-detected"
    n_slices = len(z_centers) if z_centers else "N"

    _, metric_ub, _, eq_ub = _dipole_xy_metric_sentence(orientation_metric)
    _, spatial_ubs = _dipole_xy_spatial_paragraph_text(params)

    p1 = (
        "p",
        f"The spatial distribution of water dipole orientation in the xy plane was "
        f"computed at {n_slices} discrete z-slices centred at "
        f"z_k \u2208 {{{z_centers_str}}} \u00c5. "
        f"Within each slab of thickness \u0394z = {_fmt(z_slice_width)} \u00c5, "
        "molecules whose oxygen z-coordinate satisfied |z_O \u2212 z_k| \u2264 \u0394z/2 "
        "(Eq.\u20091) were selected for analysis. "
        "A uniform xy grid of spacing "
        f"\u0394_xy = {_fmt(xy_grid_size)} \u00c5 was constructed over the full periodic box. "
        "The molecular dipole vector d was computed from the oxygen position and the mean "
        "hydrogen position (Eq.\u20092), consistent with the direction of the permanent "
        "dipole moment. " + metric_ub,
    )

    eq1 = ("eq", "(Eq.\u20091)\u2003|z_O \u2212 z_k| \u2264 \u0394z/2")
    eq2 = ("eq", "(Eq.\u20092)\u2003d = r_O \u2212 \u00bd(r_H\u2081 + r_H\u2082),\u2003d\u0302 = d / |d|")
    eq3 = ("eq", eq_ub)

    p3 = (
        "p",
        f"The trajectory was sampled at every {_fmt(step)} frame(s). For each grid cell, "
        "the final orientation value was obtained as the weighted average of all accumulated "
        "contributions (Eq.\u20094): the numerator sums the product of the spatial weight "
        "and orientation metric over all molecule\u2013frame pairs in the slab; the "
        "denominator sums the corresponding weights. Grid cells with no contributing "
        "molecules were excluded from visualisation.",
    )
    eq4 = (
        "eq",
        "(Eq.\u20094)\u2003V\u0305(x\u1d62, y\u2c7c; z\u2096) = "
        "\u03a3_{f,m} w_{f,m}\u00b7V_{f,m} / \u03a3_{f,m} w_{f,m}",
    )
    p4 = (
        "p",
        "Slices were placed near both clay\u2013water interfaces. Where slice centres were "
        "auto-detected from the clay interface boundaries, slices at the upper surface "
        "(positive-z side) were positioned at z\u2096 = z\u209a\u2091\u1d63\u1da0 \u2212 \u03b4 "
        "(water immediately below the upper clay), and slices at the lower surface "
        "(negative-z side) at z\u2096 = z\u209a\u2091\u1d63\u1da0 + \u03b4, where "
        "\u03b4 \u2208 {2, 5} \u00c5. The channel midpoint (z\u00a0=\u00a00) was also "
        "included as a reference slice. Because the dipole angle \u03b8 is always "
        "measured against the global +z axis, its directional meaning is surface-dependent: "
        "at slices near the upper clay, \u03b8\u00a0<\u00a090\u00b0 indicates preferential "
        "oxygen alignment toward that surface; at slices near the lower clay, "
        "\u03b8\u00a0>\u00a090\u00b0 signals the same orientation relative to the surface below. "
        "Consequently, maps from the two surfaces display complementary colour distributions "
        "when the colourmap is centred at 90\u00b0.",
    )

    return [p1, eq1, eq2, eq3] + spatial_ubs + [p3, eq4, p4]


# ──────────────────────────────────────────────────────────────────────────────
# TOP-LEVEL ENTRY POINT — WATER DIPOLE METHODS
# ──────────────────────────────────────────────────────────────────────────────

def write_water_dipole_methods(
    results_obj: Any = None,
    dipole_z_params: Optional[Dict[str, Any]] = None,
    dipole_xy_params: Optional[Dict[str, Any]] = None,
    output_dir: Optional[Path] = None,
    show_in_notebook: bool = True,
    save_docx: bool = True,
):
    """Generate methods text for both water dipole orientation analyses.

    Parameters
    ----------
    results_obj : ignored (kept for API consistency)
    dipole_z_params : dict with keys: bin_size, angle_bin_size
    dipole_xy_params : dict with keys: z_slice_centers, z_slice_width, xy_grid_size,
        step, orientation_metric, smoothing_method, water_vdw_radius, hydrogen_vdw_radius,
        xy_projection, projection_method, projection_radius, ellipse_aspect_ratio
    output_dir : Path for saving .docx
    show_in_notebook : bool
    save_docx : bool

    Returns
    -------
    (z_text, xy_text, saved_path)
    """
    del results_obj

    if dipole_z_params is None:
        dipole_z_params = {}
    if dipole_xy_params is None:
        dipole_xy_params = {}

    z_md = build_water_dipole_z_methods_text(dipole_z_params)
    xy_md = build_water_dipole_xy_methods_text(dipole_xy_params)

    if show_in_notebook:
        z_blocks = build_water_dipole_z_methods_manuscript_blocks(dipole_z_params)
        _show_methods_as_white_page(
            title="Methods: Water Dipole Orientation — z-Profile",
            blocks=z_blocks,
        )
        xy_blocks = build_water_dipole_xy_methods_manuscript_blocks(dipole_xy_params)
        _show_methods_as_white_page(
            title="Methods: Water Dipole Orientation — Planar XY Maps",
            blocks=xy_blocks,
        )

    saved_path = None
    if save_docx:
        out_dir = Path(output_dir) if output_dir else Path.cwd()
        out_dir.mkdir(parents=True, exist_ok=True)

        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        saved_path = out_dir / f"water_dipole_orientation_methods_{ts}.docx"

        doc = Document()

        for heading_style_name in ("Heading 1", "Heading 2"):
            if heading_style_name in doc.styles:
                doc.styles[heading_style_name].font.name = "Times New Roman"
                doc.styles[heading_style_name].font.color.rgb = RGBColor(0, 0, 0)

        for section_title, blocks in [
            (
                "Methods: Water Dipole Orientation \u2014 z-Profile",
                build_water_dipole_z_methods_manuscript_blocks(dipole_z_params),
            ),
            (
                "Methods: Water Dipole Orientation \u2014 Planar XY Maps",
                build_water_dipole_xy_methods_manuscript_blocks(dipole_xy_params),
            ),
        ]:
            heading = doc.add_heading(section_title, level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in heading.runs:
                run.font.name = "Times New Roman"
                run.font.color.rgb = RGBColor(0, 0, 0)

            for kind, text in blocks:
                line = _sanitize_xml_text(text)
                if not line:
                    continue
                p = doc.add_paragraph(line)
                if kind == "eq":
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(11)
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p.paragraph_format.space_after = Pt(8)
                    for run in p.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(11)

        doc.styles["Normal"].font.name = "Times New Roman"
        doc.styles["Normal"].font.size = Pt(12)

        doc.save(str(saved_path))
        print(f"Word file saved -> {saved_path}")

    return z_md, xy_md, saved_path


# ──────────────────────────────────────────────────────────────────────────────
# ION SPATIAL DISTRIBUTION — XY PLANAR MAPS
# ──────────────────────────────────────────────────────────────────────────────

def build_ion_spatial_xy_methods_text(params: Dict[str, Any]) -> str:
    """Build Markdown/LaTeX methods text for ion spatial distribution XY heatmaps.

    Three LD-structure paragraphs, past tense, equations inline.
    Adapts to projection_mode: 'point', 'vdw_radius', 'solvation_shell',
    'hydrated_radius'.
    """
    ion_types        = params.get("ion_types", ["NA"])
    z_slice_centers  = params.get("z_slice_centers", [])
    z_slice_width    = params.get("z_slice_width", 5.0)
    xy_grid_size     = params.get("xy_grid_size", 1.0)
    step             = params.get("step", None)
    projection_mode  = params.get("projection_mode", "vdw_radius")
    vdw_radii        = params.get("vdw_radii", {})
    solvation_radii  = params.get("solvation_radii", {})
    hydrated_radii   = params.get("hydrated_radii", {})

    ion_str = ", ".join(ion_types) if ion_types else "ions"
    z_centers_str = (
        ", ".join(_fmt(v) for v in z_slice_centers) if z_slice_centers else "auto-detected"
    )
    n_slices = len(z_slice_centers) if z_slice_centers else "N"
    step_str = f"{_fmt(step)}" if step else "all"

    # --- Projection-mode-specific paragraph ---
    if projection_mode == "point":
        proj_para = (
            r"For each ion in the slab, the ion's $xy$ position was assigned to the "
            r"nearest grid cell using a two-dimensional histogram, so that each "
            r"ion–frame observation contributed a unit count to exactly one cell. "
            r"The accumulated histogram therefore represents the total number of "
            r"ion–frame observations per cell across the trajectory."
        )
        proj_ub = (
            "For each ion in the slab, the ion's xy position was assigned to the "
            "nearest grid cell using a two-dimensional histogram, so that each "
            "ion\u2013frame observation contributed a unit count to exactly one cell. "
            "The accumulated histogram represents the total number of ion\u2013frame "
            "observations per cell across the trajectory."
        )
    elif projection_mode == "vdw_radius":
        radii_str = (
            "; ".join(f"{k}: {v}\u00a0\u00c5" for k, v in vdw_radii.items())
            if vdw_radii else "literature values"
        )
        proj_para = (
            r"Rather than treating each ion as a point, a Gaussian footprint of "
            r"width equal to the ionic van der Waals radius $r_\mathrm{vdW}$ was "
            r"distributed over the grid. For ion $i$ at position $(x_i, y_i)$, the "
            r"contribution to grid cell $(x_c, y_c)$ was",
            "",
            "$$",
            r"w(x_c, y_c) = \exp\!\left(-\frac{d^2}{2\,r_\mathrm{vdW}^2}\right), "
            r"\quad d = \sqrt{(x_c-x_i)^2+(y_c-y_i)^2}",
            "$$",
            "",
            rf"where $w$ was set to zero for cells with $w < 0.01$ (i.e., beyond "
            rf"$\approx 3\,r_\mathrm{{vdW}}$). Periodic boundary conditions were "
            rf"applied to $d$ in both lateral dimensions. "
            rf"The van der Waals radii used were: {radii_str}."
        )
        proj_ub = (
            "Rather than treating each ion as a point, a Gaussian footprint of "
            "width equal to the ionic van der Waals radius r\u1d5a\u1d48\u1d42 was "
            "distributed over the grid. For ion i at position (x\u1d62, y\u1d62), the "
            "contribution to grid cell (x\u1d9c, y\u1d9c) was w = exp(\u22120.5\u00b7d\u00b2/r\u00b2) "
            "(Eq.\u20091), where d is the periodic-image-corrected distance to the cell centre. "
            "Contributions below 1% of the peak (w < 0.01) were discarded. "
            f"The van der Waals radii used were: {radii_str}."
        )
    elif projection_mode == "solvation_shell":
        radii_str = (
            "; ".join(f"{k}: {v}\u00a0\u00c5" for k, v in solvation_radii.items())
            if solvation_radii else "literature values"
        )
        proj_para = (
            r"The projected footprint of each ion was Gaussian-weighted with the radius "
            r"of its first solvation shell $r_\mathrm{sol}$. This choice captures the "
            r"spatial extent of the ion's immediate hydration layer rather than the "
            r"bare ionic size. For ion $i$ the grid contribution was",
            "",
            "$$",
            r"w(x_c, y_c) = \exp\!\left(-\frac{d^2}{2\,r_\mathrm{sol}^2}\right)",
            "$$",
            "",
            rf"with periodic-boundary-corrected distance $d$ and a cutoff at "
            rf"$w < 0.01$. Solvation shell radii used were: {radii_str}."
        )
        proj_ub = (
            "The projected footprint of each ion was Gaussian-weighted with the radius "
            "of its first solvation shell r\u209b\u2092\u2097. For ion i the grid "
            "contribution was w = exp(\u22120.5\u00b7d\u00b2/r\u00b2) (Eq.\u20091), "
            "with periodic-boundary-corrected distance d and a cutoff at w < 0.01. "
            f"Solvation shell radii used were: {radii_str}."
        )
    else:  # hydrated_radius
        radii_str = (
            "; ".join(f"{k}: {v}\u00a0\u00c5" for k, v in hydrated_radii.items())
            if hydrated_radii else "literature values"
        )
        proj_para = (
            r"Each ion was projected onto the $xy$ grid using a Gaussian footprint "
            r"parameterised by the effective hydrated ionic radius $r_\mathrm{hyd}$, "
            r"which represents the average size of the ion together with its tightly "
            r"bound hydration shells. The contribution to grid cell $(x_c, y_c)$ was",
            "",
            "$$",
            r"w(x_c, y_c) = \exp\!\left(-\frac{d^2}{2\,r_\mathrm{hyd}^2}\right)",
            "$$",
            "",
            rf"with periodic-boundary-corrected distance $d$ and a cutoff at "
            rf"$w < 0.01$. Hydrated radii used were: {radii_str}."
        )
        proj_ub = (
            "Each ion was projected onto the xy grid using a Gaussian footprint "
            "parameterised by the effective hydrated ionic radius r\u02b0\u02b8\u1d48. "
            "The grid contribution was w = exp(\u22120.5\u00b7d\u00b2/r\u00b2) (Eq.\u20091), "
            "with periodic-boundary-corrected distance d and a cutoff at w < 0.01. "
            f"Hydrated radii used were: {radii_str}."
        )

    # Build proj paragraph lines (point mode is a plain string; others are tuples)
    if projection_mode == "point":
        proj_lines = [proj_para]
    else:
        proj_lines = list(proj_para)

    lines = [
        "## Methods: Ion Spatial Distribution in the XY Plane",
        "",
        # --- Paragraph 1: slab selection and grid ---
        r"The in-plane spatial distribution of ions was characterised at "
        rf"{n_slices} discrete $z$-slice{'s' if n_slices != 1 else ''} whose centres "
        rf"$z_k \in \{{{z_centers_str}\}}$\,Å were identified as the dominant peaks "
        r"in the one-dimensional ion number-density profile along the channel normal. "
        r"Peak positions were located by analysing the $z$-density distribution across "
        r"the trajectory and retaining peaks that exceeded a minimum height threshold and "
        r"a minimum separation distance. "
        rf"For each slab of thickness $\Delta z = {_fmt(z_slice_width)}$\,Å, "
        r"an ion at $z$-position $z_\alpha$ was included when",
        "",
        "$$",
        r"\left| z_\alpha - z_k \right| \le \frac{\Delta z}{2}",
        "$$",
        "",
        rf"A uniform $xy$ grid of spacing $\Delta_{{xy}} = {_fmt(xy_grid_size)}$\,Å "
        r"was constructed to span the full periodic box dimensions, and ion coordinates "
        r"were referenced to the box-centred frame ($z = 0$ at the channel midpoint). "
        rf"The trajectory was sampled every {step_str} frame(s), "
        r"and the accumulated grid for each z-slice was divided by the total number of "
        r"analysed frames to yield the mean ion count per cell.",
        "",
        # --- Paragraph 2: projection-mode-specific ---
    ] + proj_lines + [
        "",
        # --- Paragraph 3: top/bottom surface interpretation ---
        r"Slices were placed at both clay–water interfaces simultaneously. "
        r"A slice at positive $z_k$ samples water near the upper (positive-$z$) clay "
        r"surface, while a slice at negative $z_k$ samples water near the lower clay "
        r"surface. Because the same global coordinate system is used throughout, "
        r"comparing maps at $+z_k$ and $-z_k$ directly reveals any asymmetry in "
        r"ion adsorption between the two clay surfaces. If both surfaces are "
        r"crystallographically equivalent, the two maps are expected to display the "
        r"same spatial pattern of ion accumulation relative to the underlying clay "
        r"lattice sites.",
    ]
    return "\n".join(lines)


def build_ion_spatial_xy_methods_manuscript_blocks(params: Dict[str, Any]):
    """Build Word-friendly Unicode prose blocks for ion spatial XY methods."""
    ion_types        = params.get("ion_types", ["NA"])
    z_slice_centers  = params.get("z_slice_centers", [])
    z_slice_width    = params.get("z_slice_width", 5.0)
    xy_grid_size     = params.get("xy_grid_size", 1.0)
    step             = params.get("step", None)
    projection_mode  = params.get("projection_mode", "vdw_radius")
    vdw_radii        = params.get("vdw_radii", {})
    solvation_radii  = params.get("solvation_radii", {})
    hydrated_radii   = params.get("hydrated_radii", {})

    ion_str = ", ".join(ion_types) if ion_types else "ions"
    z_centers_str = (
        ", ".join(_fmt(v) for v in z_slice_centers) if z_slice_centers else "auto-detected"
    )
    n_slices = len(z_slice_centers) if z_slice_centers else "N"
    step_str = f"{_fmt(step)}" if step else "all"

    p1 = (
        "p",
        f"The in-plane spatial distribution of ions was characterised at "
        f"{n_slices} discrete z-slice{'s' if n_slices != 1 else ''} whose centres "
        f"z\u2096 \u2208 {{{z_centers_str}}} \u00c5 were identified as the dominant "
        "peaks in the one-dimensional ion number-density profile along the channel "
        "normal. Peak positions were located by analysing the z-density distribution "
        "across the trajectory and retaining peaks that exceeded a minimum height "
        "threshold and a minimum separation distance (Eq.\u20091). "
        f"For each slab of thickness \u0394z = {_fmt(z_slice_width)} \u00c5, "
        "an ion was included when |z\u03b1 \u2212 z\u2096| \u2264 \u0394z/2. "
        f"A uniform xy grid of spacing \u0394_xy = {_fmt(xy_grid_size)} \u00c5 "
        "was constructed over the full periodic box. Ion coordinates were referenced "
        "to the box-centred frame (z\u00a0=\u00a00 at the channel midpoint). "
        f"The trajectory was sampled every {step_str} frame(s), and each accumulated "
        "grid was divided by the number of analysed frames to yield the mean ion "
        "count per cell.",
    )
    eq1 = ("eq", "(Eq.\u20091)\u2003|z\u03b1 \u2212 z\u2096| \u2264 \u0394z/2")

    if projection_mode == "point":
        p2 = (
            "p",
            "For each ion in the slab, the ion's xy position was assigned to the "
            "nearest grid cell using a two-dimensional histogram, so that each "
            "ion\u2013frame observation contributed a unit count to exactly one cell. "
            "The accumulated histogram represents the total number of ion\u2013frame "
            "observations per cell across the trajectory.",
        )
        proj_blocks = [p2]
    else:
        if projection_mode == "vdw_radius":
            radii_dict = vdw_radii
            radius_label = "van der Waals"
            r_sym = "r\u1d5a\u1d48\u1d42"
        elif projection_mode == "solvation_shell":
            radii_dict = solvation_radii
            radius_label = "first solvation shell"
            r_sym = "r\u209b\u2092\u2097"
        else:
            radii_dict = hydrated_radii
            radius_label = "hydrated ionic"
            r_sym = "r\u02b0\u02b8\u1d48"
        radii_str = (
            "; ".join(f"{k}: {v}\u00a0\u00c5" for k, v in radii_dict.items())
            if radii_dict else "literature values"
        )
        p2 = (
            "p",
            f"Rather than treating each ion as a point, a Gaussian footprint parameterised "
            f"by the ionic {radius_label} radius {r_sym} was distributed over the grid. "
            f"For ion i at position (x\u1d62, y\u1d62), the contribution to grid cell "
            "(x\u1d9c, y\u1d9c) was w = exp(\u22120.5\u00b7d\u00b2/r\u00b2) (Eq.\u20092), "
            "where d is the periodic-image-corrected lateral distance. "
            "Contributions below 1% of the peak (w < 0.01) were discarded. "
            f"{radius_label.capitalize()} radii used: {radii_str}.",
        )
        eq2 = (
            "eq",
            "(Eq.\u20092)\u2003w(x\u1d9c, y\u1d9c) = exp(\u22120.5\u00b7d\u00b2/"
            + r_sym + "\u00b2),\u2003d = \u221a[(x\u1d9c\u2212x\u1d62)\u00b2+(y\u1d9c\u2212y\u1d62)\u00b2]",
        )
        proj_blocks = [p2, eq2]

    p3 = (
        "p",
        "Slices were placed at both clay\u2013water interfaces simultaneously. "
        "A slice at positive z\u2096 samples ions near the upper (positive-z) clay "
        "surface, while a slice at negative z\u2096 samples ions near the lower clay "
        "surface. Because the same global coordinate system is used throughout, "
        "comparing maps at +z\u2096 and \u2212z\u2096 directly reveals any asymmetry in "
        "ion adsorption between the two clay surfaces. If both surfaces are "
        "crystallographically equivalent, the two maps display the same spatial "
        "pattern of ion accumulation relative to the underlying clay lattice sites.",
    )

    return [p1, eq1] + proj_blocks + [p3]


def write_ion_spatial_xy_methods(
    results_obj: Any = None,
    params: Optional[Dict[str, Any]] = None,
    output_dir: Optional[Path] = None,
    show_in_notebook: bool = True,
    save_docx: bool = True,
):
    """Generate methods text for the ion spatial distribution XY analysis.

    Parameters
    ----------
    results_obj : ignored (kept for API consistency)
    params : dict with keys: ion_types, z_slice_centers, z_slice_width,
        xy_grid_size, step, projection_mode, vdw_radii, solvation_radii,
        hydrated_radii
    output_dir : Path for saving .docx
    show_in_notebook : bool
    save_docx : bool

    Returns
    -------
    (text, saved_path)
    """
    del results_obj
    if params is None:
        params = {}

    md = build_ion_spatial_xy_methods_text(params)

    if show_in_notebook:
        blocks = build_ion_spatial_xy_methods_manuscript_blocks(params)
        _show_methods_as_white_page(
            title="Methods: Ion Spatial Distribution — Planar XY Maps",
            blocks=blocks,
        )

    saved_path = None
    if save_docx:
        out_dir = Path(output_dir) if output_dir else Path.cwd()
        out_dir.mkdir(parents=True, exist_ok=True)

        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        saved_path = out_dir / f"ion_spatial_xy_methods_{ts}.docx"

        doc = Document()

        for heading_style_name in ("Heading 1", "Heading 2"):
            if heading_style_name in doc.styles:
                doc.styles[heading_style_name].font.name = "Times New Roman"
                doc.styles[heading_style_name].font.color.rgb = RGBColor(0, 0, 0)

        heading = doc.add_heading(
            "Methods: Ion Spatial Distribution \u2014 Planar XY Maps", level=1
        )
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in heading.runs:
            run.font.name = "Times New Roman"
            run.font.color.rgb = RGBColor(0, 0, 0)

        for kind, text in build_ion_spatial_xy_methods_manuscript_blocks(params):
            line = _sanitize_xml_text(text)
            if not line:
                continue
            p = doc.add_paragraph(line)
            if kind == "eq":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.space_after = Pt(8)
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)

        doc.styles["Normal"].font.name = "Times New Roman"
        doc.styles["Normal"].font.size = Pt(12)

        doc.save(str(saved_path))
        print(f"Word file saved -> {saved_path}")

    return md, saved_path


# ──────────────────────────────────────────────────────────────────────────────
# ELECTROSTATIC POTENTIAL — XY PLANAR MAPS + 3D HISTOGRAM SURFACE
# ──────────────────────────────────────────────────────────────────────────────

def build_electrostatic_potential_methods_text(params: Dict[str, Any]) -> str:
    """Build Markdown/LaTeX methods text for electrostatic potential XY maps + 3D histogram.

    Three LD-structure paragraphs, past tense, equations referenced inline.
    """
    z_centers = params.get("z_slice_centers", [])
    z_slice_width = params.get("z_slice_width", 2.0)
    xy_grid_size = params.get("xy_grid_size", 0.5)
    step = params.get("step", 1)
    potential_calculation_method = params.get("potential_calculation_method", "point_charges")
    dielectric_screening = params.get("dielectric_screening", True)
    relative_permittivity = params.get("relative_permittivity", 78.0)
    cutoff_distance = params.get("cutoff_distance", 15.0)
    gaussian_smoothing = params.get("gaussian_smoothing", True)
    smoothing_sigma = params.get("smoothing_sigma", None)
    if smoothing_sigma is None:
        smoothing_sigma = xy_grid_size
    return_negative_potential = params.get("return_negative_potential", True)
    surface_peak_method = params.get("surface_peak_method", "global_max")
    return_surface_peaks = params.get("return_surface_peaks", False)

    z_centers_str = ", ".join(_fmt(v) for v in z_centers) if z_centers else "auto-detected"
    n_slices = len(z_centers) if z_centers else "N"
    step_str = _fmt(step) if step else "1"

    # screening sentence
    screening_str = (
        rf", after dividing all pairwise contributions by the relative permittivity "
        rf"$\varepsilon_r = {_fmt(relative_permittivity)}$ of bulk water "
        r"to account for dielectric screening in the aqueous medium"
        if dielectric_screening else ""
    )

    # smoothing sentence
    smoothing_str = (
        rf" The potential grids were subsequently smoothed with an isotropic "
        rf"Gaussian kernel of width $\sigma = {_fmt(smoothing_sigma)}$ Å to suppress "
        r"fluctuations arising from finite ion sampling."
        if gaussian_smoothing else ""
    )

    # sign convention sentence
    if return_negative_potential:
        sign_str = (
            r" For visualisation, the sign-inverted potential "
            r"$\Phi = -\langle\varphi\rangle$ was used so that regions of strong "
            r"negative electrostatic potential (attracting cations) appear as peaks "
            r"of large positive height."
        )
    else:
        sign_str = ""

    # surface peak sentence
    if return_surface_peaks:
        peak_method_str = {
            "global_max": "the global maximum across all $z$-slices",
            "weighted_average": "a frame-count-weighted average across $z$-slices",
        }.get(surface_peak_method, f"the '{surface_peak_method}' method")
        peak_str = (
            rf" A memory-efficient surface representation was derived by taking, "
            rf"at each $(x_i, y_j)$ position, {peak_method_str} of "
            rf"$\Phi(x_i, y_j;\, z_k)$ to yield a single projected surface "
            r"$\Phi_\mathrm{surf}(x_i, y_j)$."
        )
    else:
        peak_str = ""

    lines = [
        "## Methods: Electrostatic Potential Spatial Distribution",
        "",
        # --- Paragraph 1: grid geometry and ion slab assignment ---
        rf"The lateral distribution of the electrostatic potential generated by dissolved "
        rf"ions was mapped in the $xy$ plane at {n_slices} discrete $z$-positions. "
        rf"Each $z$-slice of thickness $\Delta z = {_fmt(z_slice_width)}$ Å was centred at "
        rf"$z_k \in \{{{z_centers_str}\}}$ Å, and ions falling within $|z \!-\! z_k| "
        r"\le \Delta z / 2$ were included in that slice. "
        rf"A uniform two-dimensional grid with spacing $\Delta_{{xy}} = {_fmt(xy_grid_size)}$ Å "
        r"was constructed over the full periodic $xy$ box for each slice, and the trajectory "
        rf"was sampled every {step_str} frame(s).",
        "",
        # --- Paragraph 2: Coulomb summation, screening, smoothing, averaging ---
        r"For each sampled frame, the electrostatic potential $\varphi(\mathbf{r}_g;\, z_k)$ "
        r"at every grid point $\mathbf{r}_g = (x_i, y_j)$ was evaluated as a direct "
        r"Coulomb sum over all ions within the cutoff "
        rf"$r_\mathrm{{cut}} = {_fmt(cutoff_distance)}$ Å{screening_str}:",
        "",
        "$$",
        r"\varphi(x_i, y_j;\, z_k) = \sum_{n:\,r_{n,g} \le r_{\mathrm{cut}}} "
        r"\frac{q_n}{r_{n,g}}",
        "$$",
        "",
        r"where $q_n$ is the partial charge of ion $n$ and $r_{n,g}$ is the "
        r"minimum-image distance from $n$ to grid point $g$. "
        r"Frame-level grids were accumulated and normalised by the total number of "
        r"sampled frames to give the time-averaged potential "
        r"$\langle\varphi\rangle(x_i, y_j;\, z_k)$."
        + smoothing_str
        + sign_str,
        "",
        # --- Paragraph 3: 2D maps, 3D histogram bars, surface projection ---
        r"The time-averaged potential was visualised in two complementary ways. "
        r"First, per-slice two-dimensional heatmaps of "
        r"$\Phi(x_i, y_j;\, z_k)$ were produced to reveal the lateral pattern of "
        r"electrostatic attraction and repulsion at each selected $z$-plane. "
        r"Second, three-dimensional histogram surfaces were constructed in which the "
        r"height of each bar is proportional to $\Phi$ at that grid position, with "
        r"bars coloured by potential magnitude; the clay framework was overlaid as "
        r"isocontour lines of the Si and MgOH atom density. "
        + peak_str
        + r" Together, the 2D maps and 3D histogram surfaces show how the "
        r"electrostatic landscape is modulated by the periodic arrangement of the "
        r"ditrigonal cavities in the clay tetrahedral sheet.",
    ]

    return "\n".join(lines)


def build_electrostatic_potential_methods_manuscript_blocks(params: Dict[str, Any]):
    """Build Word-friendly Unicode prose blocks for the electrostatic potential methods.

    Three LD-structure paragraphs with equations referenced inline, past tense.
    """
    z_centers = params.get("z_slice_centers", [])
    z_slice_width = params.get("z_slice_width", 2.0)
    xy_grid_size = params.get("xy_grid_size", 0.5)
    step = params.get("step", 1)
    potential_calculation_method = params.get("potential_calculation_method", "point_charges")
    dielectric_screening = params.get("dielectric_screening", True)
    relative_permittivity = params.get("relative_permittivity", 78.0)
    cutoff_distance = params.get("cutoff_distance", 15.0)
    gaussian_smoothing = params.get("gaussian_smoothing", True)
    smoothing_sigma = params.get("smoothing_sigma", None)
    if smoothing_sigma is None:
        smoothing_sigma = xy_grid_size
    return_negative_potential = params.get("return_negative_potential", True)
    surface_peak_method = params.get("surface_peak_method", "global_max")
    return_surface_peaks = params.get("return_surface_peaks", False)

    z_centers_str = ", ".join(_fmt(v) for v in z_centers) if z_centers else "auto-detected"
    n_slices = len(z_centers) if z_centers else "N"
    step_str = _fmt(step) if step else "1"

    screening_ub = (
        f", after dividing all pairwise contributions by the relative permittivity "
        f"\u03b5_r\u2009=\u2009{_fmt(relative_permittivity)} of bulk water "
        "to account for dielectric screening"
        if dielectric_screening else ""
    )
    smoothing_ub = (
        f" Each potential grid was subsequently smoothed with an isotropic Gaussian "
        f"kernel of width \u03c3\u2009=\u2009{_fmt(smoothing_sigma)}\u2009\u00c5 to "
        "suppress sampling fluctuations."
        if gaussian_smoothing else ""
    )
    sign_ub = (
        " For visualisation the sign-inverted potential \u03a6 = \u2212\u27e8\u03c6\u27e9 "
        "was used so that regions of strong negative electrostatic potential "
        "(attracting cations) appear as high-magnitude peaks."
        if return_negative_potential else ""
    )
    if return_surface_peaks:
        peak_method_str = {
            "global_max": "the global maximum across all z-slices",
            "weighted_average": "a frame-count-weighted average across z-slices",
        }.get(surface_peak_method, f"the '{surface_peak_method}' method")
        peak_ub = (
            f" A projected surface \u03a6_surf(x\u1d62, y\u2c7c) was derived by taking "
            f"{peak_method_str} of \u03a6 at each xy position."
        )
    else:
        peak_ub = ""

    return [
        # --- Paragraph 1 ---
        (
            "p",
            f"The lateral distribution of the electrostatic potential generated by dissolved "
            f"ions was mapped in the xy plane at {n_slices} discrete z-positions. "
            f"Each z-slice of thickness \u0394z\u2009=\u2009{_fmt(z_slice_width)}\u2009\u00c5 "
            f"was centred at z_k \u2208 {{{z_centers_str}}}\u2009\u00c5, and ions falling within "
            f"|\u2009z \u2212 z_k\u2009| \u2264 \u0394z/2 were included in that slice. "
            f"A uniform xy grid with spacing \u0394_xy\u2009=\u2009{_fmt(xy_grid_size)}\u2009\u00c5 "
            f"was constructed over the full periodic box, and the trajectory was sampled every "
            f"{step_str} frame(s).",
        ),
        # --- Equation 1 ---
        (
            "eq",
            "(Eq.\u20091)\u2003\u03c6(x\u1d62, y\u2c7c;\u2003z_k) = "
            "\u2211_n  q_n / r_{n,g}    (sum over ions with r_{n,g} \u2264 r_cut)",
        ),
        # --- Paragraph 2 ---
        (
            "p",
            f"For each sampled frame, the electrostatic potential \u03c6(r_g;\u2003z_k) at every "
            f"grid point r_g\u2009=\u2009(x\u1d62, y\u2c7c) was evaluated as a direct Coulomb sum "
            f"(Eq.\u20091) over all ions within the cutoff r_cut\u2009=\u2009{_fmt(cutoff_distance)}\u2009\u00c5"
            f"{screening_ub}. "
            f"Frame-level grids were accumulated and normalised by the total number of sampled "
            f"frames to give the time-averaged potential \u27e8\u03c6\u27e9(x\u1d62, y\u2c7c;\u2003z_k)."
            + smoothing_ub
            + sign_ub,
        ),
        # --- Paragraph 3 ---
        (
            "p",
            "The time-averaged potential was visualised in two complementary ways. "
            "First, per-slice two-dimensional heatmaps of \u03a6(x\u1d62, y\u2c7c;\u2003z_k) revealed "
            "the lateral pattern of electrostatic attraction and repulsion at each selected z-plane. "
            "Second, three-dimensional histogram surfaces were constructed in which the height of "
            "each bar is proportional to \u03a6 at that grid position, with bars coloured by "
            "potential magnitude; the clay framework was overlaid as isocontour lines of the "
            "Si and MgOH atom density."
            + peak_ub
            + " Together, the 2D maps and 3D histogram surfaces show how the electrostatic "
            "landscape is modulated by the periodic arrangement of the ditrigonal cavities in "
            "the clay tetrahedral sheet.",
        ),
    ]


def write_electrostatic_potential_methods(
    results_obj=None,
    params: Optional[Dict[str, Any]] = None,
    output_dir: Optional[Path] = None,
    show_in_notebook: bool = True,
    save_docx: bool = True,
):
    """Convenience wrapper: write electrostatic potential methods for notebook use.

    Parameters
    ----------
    results_obj : ignored (accepted for API symmetry)
    params : dict
        Keys drawn from ``calculate_electrostatic_potential_spatial_distribution_xy``
        and ``plot_electrostatic_potential_3d_histograms_flipped3`` parameters.
        Commonly used keys: z_slice_centers, z_slice_width, xy_grid_size, step,
        potential_calculation_method, dielectric_screening, relative_permittivity,
        cutoff_distance, gaussian_smoothing, smoothing_sigma, return_negative_potential,
        return_surface_peaks, surface_peak_method.
    output_dir : Path-like, optional
    show_in_notebook : bool
    save_docx : bool

    Returns
    -------
    (text, saved_path)
    """
    del results_obj
    if params is None:
        params = {}

    md = build_electrostatic_potential_methods_text(params)

    if show_in_notebook:
        blocks = build_electrostatic_potential_methods_manuscript_blocks(params)
        _show_methods_as_white_page(
            title="Methods: Electrostatic Potential Spatial Distribution",
            blocks=blocks,
        )

    saved_path = None
    if save_docx:
        out_dir = Path(output_dir) if output_dir else Path.cwd()
        out_dir.mkdir(parents=True, exist_ok=True)

        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        saved_path = out_dir / f"electrostatic_potential_methods_{ts}.docx"

        doc = Document()

        for heading_style_name in ("Heading 1", "Heading 2"):
            if heading_style_name in doc.styles:
                doc.styles[heading_style_name].font.name = "Times New Roman"
                doc.styles[heading_style_name].font.color.rgb = RGBColor(0, 0, 0)

        heading = doc.add_heading(
            "Methods: Electrostatic Potential Spatial Distribution", level=1
        )
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in heading.runs:
            run.font.name = "Times New Roman"
            run.font.color.rgb = RGBColor(0, 0, 0)

        for kind, text in build_electrostatic_potential_methods_manuscript_blocks(params):
            line = _sanitize_xml_text(text)
            if not line:
                continue
            p = doc.add_paragraph(line)
            if kind == "eq":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.space_after = Pt(8)
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)

        doc.styles["Normal"].font.name = "Times New Roman"
        doc.styles["Normal"].font.size = Pt(12)

        doc.save(str(saved_path))
        print(f"Word file saved -> {saved_path}")

    return md, saved_path


# ──────────────────────────────────────────────────────────────────────────────
# WATER COORDINATION / SOLVATION SHELLS VS Z + 3D SURFACE VISUALISATION
# ──────────────────────────────────────────────────────────────────────────────

def build_coordination_shells_methods_text(params: Dict[str, Any]) -> str:
    """Build Markdown/LaTeX methods text for solvation-shell coordination analysis + 3D surfaces."""

    # ---- calculation params ----
    shell_radii = params.get("shell_radii", [3.5, 6.0, 9.0])
    ion_specific_radii = params.get("ion_specific_radii", {})
    step = params.get("step", 1)
    target_z_values = params.get("target_z_values", [])
    z_slice_width = params.get("z_slice_width", 1.0)
    xy_grid_size = params.get("xy_grid_size", 1.0)
    ion_types = params.get("ion_types", [])

    # ---- surface creation params ----
    smoothing = params.get("smoothing", True)
    smoothing_sigma = params.get("smoothing_sigma", 1.0)
    z_scale_factor = params.get("z_scale_factor", 1.0)
    shell_selection = params.get("shell_selection", ["first_shell", "second_shell", "third_shell"])

    # ---- helpers ----
    def _radii_str(radii):
        return ", ".join(_fmt(r) for r in radii)

    ion_str = ", ".join(ion_types) if ion_types else "all ions"
    z_centers_str = ", ".join(_fmt(v) for v in target_z_values) if target_z_values else "auto-detected"
    n_slices = len(target_z_values) if target_z_values else "N"
    n_shells = len(shell_selection)

    # default radii sentence
    default_radii_str = _radii_str(shell_radii)
    shell_bounds = []
    for i, r in enumerate(shell_radii):
        lo = _fmt(shell_radii[i - 1]) if i > 0 else "0"
        shell_bounds.append(rf"shell {i+1}: {lo}–{_fmt(r)} Å")

    # ion-specific radii sentences
    ion_radii_lines = []
    for ion, radii in ion_specific_radii.items():
        bounds = []
        for i, r in enumerate(radii):
            lo = _fmt(radii[i - 1]) if i > 0 else "0"
            bounds.append(rf"{lo}–{_fmt(r)} Å")
        ion_radii_lines.append(rf"  - {ion}: {'; '.join(bounds)}")
    ion_radii_block = ("\n".join(ion_radii_lines) + "\n") if ion_radii_lines else ""

    smoothing_str = (
        rf" A Gaussian kernel of width $\sigma = {_fmt(smoothing_sigma)}$ (in grid units) "
        r"was applied to each coordination surface to suppress single-frame sampling noise."
        if smoothing else ""
    )

    shell_names_latex = {"first_shell": r"first", "second_shell": r"second", "third_shell": r"third"}
    shell_list_str = ", ".join(shell_names_latex.get(s, s) for s in shell_selection)

    lines = [
        "## Methods: Ion Hydration Shells vs. z — Spatial Analysis and 3D Coordination Surfaces",
        "",
        # ---- Para 1: radial shell definitions ----
        rf"The hydration structure around dissolved ions ({ion_str}) was quantified as a "
        rf"function of the lateral position and $z$-distance from the clay surface. "
        rf"Solvation shells were defined by concentric radial cutoffs: "
        + "; ".join(shell_bounds)
        + r" from each ion centre. "
        + (
            r"Ion-specific radii were used where applicable: "
            + "; ".join(
                f"{ion}: " + "/".join(_fmt(r) for r in radii) + " Å"
                for ion, radii in ion_specific_radii.items()
            )
            + r" (in order, first–third shell)."
            if ion_specific_radii else ""
        ),
        "",
        # ---- Para 2: z-layer sampling + water counting ----
        rf"The trajectory was sampled every {_fmt(step)} frame(s) and, at each frame, "
        rf"a $z$-slab of thickness $\Delta z = {_fmt(z_slice_width)}$ Å centred at "
        rf"$z_k \in \{{{z_centers_str}\}}$ Å was interrogated. "
        r"Ions falling within $|z - z_k| \le \Delta z/2$ were identified and, for each "
        r"such ion, the number of water oxygen atoms (OW) within each concentric radial "
        r"shell was counted using the minimum-image distance metric. "
        r"The per-ion coordination numbers $\bar{n}_\mathrm{shell}$ were then "
        r"projected onto a uniform two-dimensional grid "
        rf"(spacing $\Delta_{{xy}} = {_fmt(xy_grid_size)}$ Å, "
        r"dimensions set by the $xy$ simulation box) by binning each ion's lateral "
        r"position $(x, y)$ — giving a mean coordination-number map "
        r"$\langle n_\mathrm{shell} \rangle(x_i, y_j;\, z_k)$ for each slice and shell.",
        "",
        # ---- Para 3: height surfaces + visualisation ----
        rf"Spatially resolved coordination maps were converted to "
        rf"three-dimensional height surfaces in which the surface elevation at "
        r"$(x_i, y_j)$ directly encodes the local coordination number, scaled by "
        rf"a factor of {_fmt(z_scale_factor)}."
        + smoothing_str
        + rf" Surfaces for the {shell_list_str} shells "
        r"were visualised using distinct colour maps per shell, revealing "
        r"how the clay ditrigonal cavity pattern modulates the hydration "
        r"structure of near-surface ions.",
    ]

    return "\n".join(lines)


def build_coordination_shells_methods_manuscript_blocks(params: Dict[str, Any]):
    """Build Word-friendly Unicode blocks for coordination-shell methods."""

    shell_radii = params.get("shell_radii", [3.5, 6.0, 9.0])
    ion_specific_radii = params.get("ion_specific_radii", {})
    step = params.get("step", 1)
    target_z_values = params.get("target_z_values", [])
    z_slice_width = params.get("z_slice_width", 1.0)
    xy_grid_size = params.get("xy_grid_size", 1.0)
    ion_types = params.get("ion_types", [])
    smoothing = params.get("smoothing", True)
    smoothing_sigma = params.get("smoothing_sigma", 1.0)
    z_scale_factor = params.get("z_scale_factor", 1.0)
    shell_selection = params.get("shell_selection", ["first_shell", "second_shell", "third_shell"])

    ion_str = ", ".join(ion_types) if ion_types else "all ions"
    z_centers_str = ", ".join(_fmt(v) for v in target_z_values) if target_z_values else "auto-detected"

    shell_bounds = []
    for i, r in enumerate(shell_radii):
        lo = _fmt(shell_radii[i - 1]) if i > 0 else "0"
        shell_bounds.append(f"shell {i+1}: {lo}\u2013{_fmt(r)}\u2009\u00c5")

    ion_radii_bits = []
    for ion, radii in ion_specific_radii.items():
        ion_radii_bits.append(f"{ion}: " + "/".join(_fmt(r) for r in radii) + "\u2009\u00c5")

    smoothing_ub = (
        f" A Gaussian kernel of width \u03c3\u2009=\u2009{_fmt(smoothing_sigma)} (grid units) "
        "was applied to each coordination surface to suppress sampling noise."
        if smoothing else ""
    )

    shell_name_map = {"first_shell": "first", "second_shell": "second", "third_shell": "third"}
    shell_list_str = ", ".join(shell_name_map.get(s, s) for s in shell_selection)

    return [
        (
            "p",
            f"The hydration structure around dissolved ions ({ion_str}) was quantified as a "
            f"function of lateral position and z-distance from the clay surface. "
            f"Solvation shells were defined by concentric radial cutoffs: "
            + "; ".join(shell_bounds)
            + "."
            + (
                " Ion-specific radii were applied where available: "
                + "; ".join(ion_radii_bits)
                + " (first through third shell)."
                if ion_radii_bits else ""
            ),
        ),
        (
            "p",
            f"The trajectory was sampled every {_fmt(step)} frame(s). At each frame, "
            f"a z-slab of thickness \u0394z\u2009=\u2009{_fmt(z_slice_width)}\u2009\u00c5 centred at "
            f"z_k \u2208 {{{z_centers_str}}}\u2009\u00c5 was interrogated. "
            f"For each ion within the slab, the number of water oxygen (OW) atoms "
            f"within each radial shell was counted using the minimum-image convention. "
            f"Per-ion coordination numbers were binned onto a uniform xy grid "
            f"(\u0394_xy\u2009=\u2009{_fmt(xy_grid_size)}\u2009\u00c5) by each ion\u2019s "
            f"lateral (x, y) position, yielding a mean coordination-number map "
            f"\u27e8n_shell\u27e9(x\u1d62, y\u2c7c;\u2003z_k) per slice and shell.",
        ),
        (
            "p",
            f"Spatially resolved coordination maps were rendered as three-dimensional "
            f"height surfaces in which elevation encodes local coordination number "
            f"(scale factor {_fmt(z_scale_factor)})."
            + smoothing_ub
            + f" Surfaces for the {shell_list_str} shells were displayed with "
            f"per-shell colour maps, revealing how the clay ditrigonal cavity pattern "
            f"modulates the hydration structure of near-surface ions.",
        ),
    ]


def write_coordination_shells_methods(
    results_obj=None,
    params: Optional[Dict[str, Any]] = None,
    output_dir: Optional[Path] = None,
    show_in_notebook: bool = True,
    save_docx: bool = True,
):
    """Convenience wrapper: write coordination-shell methods for notebook use.

    Parameters
    ----------
    results_obj : ignored
    params : dict
        Keys drawn from ``calculate_solvation_shells_vs_z_detailed`` and
        ``create_coordination_height_surfaces`` / ``plot_coordination_height_surfaces``.
        Recommended keys: shell_radii, ion_specific_radii, step, target_z_values,
        z_slice_width, xy_grid_size, ion_types, smoothing, smoothing_sigma,
        z_scale_factor, shell_selection.
    output_dir : Path-like, optional
    show_in_notebook : bool
    save_docx : bool

    Returns
    -------
    (text, saved_path)
    """
    del results_obj
    if params is None:
        params = {}

    md = build_coordination_shells_methods_text(params)

    if show_in_notebook:
        blocks = build_coordination_shells_methods_manuscript_blocks(params)
        _show_methods_as_white_page(
            title="Methods: Ion Hydration Shells \u2014 Spatial Analysis and 3D Surfaces",
            blocks=blocks,
        )

    saved_path = None
    if save_docx:
        out_dir = Path(output_dir) if output_dir else Path.cwd()
        out_dir.mkdir(parents=True, exist_ok=True)

        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        saved_path = out_dir / f"coordination_shells_methods_{ts}.docx"

        doc = Document()

        for heading_style_name in ("Heading 1", "Heading 2"):
            if heading_style_name in doc.styles:
                doc.styles[heading_style_name].font.name = "Times New Roman"
                doc.styles[heading_style_name].font.color.rgb = RGBColor(0, 0, 0)

        heading = doc.add_heading(
            "Methods: Ion Hydration Shells \u2014 Spatial Analysis and 3D Surfaces", level=1
        )
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in heading.runs:
            run.font.name = "Times New Roman"
            run.font.color.rgb = RGBColor(0, 0, 0)

        for kind, text in build_coordination_shells_methods_manuscript_blocks(params):
            line = _sanitize_xml_text(text)
            if not line:
                continue
            p = doc.add_paragraph(line)
            if kind == "eq":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.space_after = Pt(8)
            for run in p.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)

        doc.styles["Normal"].font.name = "Times New Roman"
        doc.styles["Normal"].font.size = Pt(12)

        doc.save(str(saved_path))
        print(f"Word file saved -> {saved_path}")

    return md, saved_path


# ──────────────────────────────────────────────────────────────────────────────
# RDF vs Z-POSITION — ION–WATER AND ION–ION PAIRS
# ──────────────────────────────────────────────────────────────────────────────

def build_rdf_vs_z_methods_text(params: Dict[str, Any]) -> str:
    """Build Markdown/LaTeX methods text for g(r,z) RDF analysis."""

    ion_pairs        = params.get("ion_pairs", [])
    r_max            = params.get("r_max", 12.0)
    dr               = params.get("dr", 0.1)
    z_bin_width      = params.get("z_bin_width", 1.0)
    step             = params.get("step", 1)
    density_mode     = params.get("density_mode", "global")
    ion_pair_radii   = params.get("ion_pair_radii", {})   # pairing boundaries dict

    # Classify pairs
    water_pairs = [p for p in ion_pairs if p.split("-")[-1].upper() in ("OW", "O", "WATER")]
    ion_ion_pairs = [p for p in ion_pairs if p not in water_pairs]

    water_str    = ", ".join(water_pairs)   if water_pairs    else "none"
    ion_ion_str  = ", ".join(ion_ion_pairs) if ion_ion_pairs  else "none"

    density_desc = {
        "global":       (r"bulk density of the simulation cell (total water molecules "
                         r"divided by total cell volume)"),
        "local":        r"local z-slice density (water molecules in each slab divided by slab volume)",
        "water_region": r"density of the water-containing region only",
    }.get(density_mode, rf"'{density_mode}' density")

    # Ion-pairing boundary sentence
    pairing_str = ""
    if ion_pair_radii:
        pairs_listed = "; ".join(
            f"{pair}: " + "/".join(_fmt(r) for r in radii) + " Å"
            for pair, radii in ion_pair_radii.items()
        )
        pairing_str = (
            r" Boundaries separating contact ion pairs (CIP), "
            r"solvent-shared ion pairs (SIP), and doubly solvent-separated "
            r"ion pairs (DSIP) were drawn at: "
            + pairs_listed + r"."
        )

    n_r = int(round(r_max / dr))

    lines = [
        "## Methods: Radial Distribution Functions vs. z-Position",
        "",
        # --- Para 1: g(r,z) definition and z-slice setup ---
        rf"Radial distribution functions (RDFs) were calculated as a function of "
        rf"$z$-distance from the clay surface to characterise the local solvation "
        rf"structure and ion-pairing state of dissolved species. "
        rf"The simulation box was divided into $z$-slabs of thickness "
        rf"$\Delta z = {_fmt(z_bin_width)}$ Å, and the trajectory was sampled every "
        rf"{_fmt(step)} frame(s). "
        rf"For each reference ion found within a given slab, distances to all target "
        rf"atoms in the entire simulation box were computed using the minimum-image "
        rf"convention with full periodic boundary conditions, retaining distances "
        rf"$0 < r \le {_fmt(r_max)}$ Å in {n_r} bins of width $\Delta r = {_fmt(dr)}$ Å.",
        "",
        # --- Para 2: normalisation ---
        r"The pair correlation function was obtained from the standard normalisation",
        "",
        "$$",
        r"g(r;\,z_k) = \frac{\langle n(r,\Delta r;\,z_k)\rangle / N_\mathrm{ref}(z_k)}"
        r"{\rho_\mathrm{ref} \cdot 4\pi r^2 \Delta r}",
        "$$",
        "",
        rf"where $\langle n(r,\Delta r;\,z_k)\rangle$ is the mean number of target atoms "
        rf"found in the shell $[r, r+\Delta r)$ around a reference ion in slab $k$, "
        rf"$N_{{\mathrm{{ref}}}}(z_k)$ is the number of reference ions in that slab, and "
        rf"$\rho_{{\mathrm{{ref}}}}$ is the {density_desc}. "
        rf"Ion–ion pairs always used the local z-slice density of the target species. "
        rf"Pairs analysed: ion–water ({water_str}); ion–ion ({ion_ion_str}).",
        "",
        # --- Para 3: ion-pairing classification (only if boundaries given) ---
        *(
            [
                r"For cation–anion pairs, the $g(r;\,z)$ profiles were interpreted "
                r"in terms of ion-pairing speciation. "
                r"Contact ion pairs (CIP) correspond to direct cation–anion contact, "
                r"solvent-shared ion pairs (SIP) to a pair separated by a single "
                r"water molecule, and doubly solvent-separated ion pairs (DSIP) to "
                r"pairs separated by two solvent layers."
                + pairing_str
                + r" The $g(r;\,z)$ heat maps were displayed with boundaries shown as "
                r"dashed vertical lines to indicate changes in pairing state as a "
                r"function of proximity to the clay surface.",
            ]
            if ion_ion_pairs else []
        ),
    ]

    return "\n".join(lines)


def build_rdf_vs_z_methods_manuscript_blocks(params: Dict[str, Any]):
    """Build Word-friendly Unicode blocks for g(r,z) RDF methods."""

    ion_pairs        = params.get("ion_pairs", [])
    r_max            = params.get("r_max", 12.0)
    dr               = params.get("dr", 0.1)
    z_bin_width      = params.get("z_bin_width", 1.0)
    step             = params.get("step", 1)
    density_mode     = params.get("density_mode", "global")
    ion_pair_radii   = params.get("ion_pair_radii", {})

    water_pairs   = [p for p in ion_pairs if p.split("-")[-1].upper() in ("OW","O","WATER")]
    ion_ion_pairs = [p for p in ion_pairs if p not in water_pairs]
    water_str    = ", ".join(water_pairs)   if water_pairs    else "none"
    ion_ion_str  = ", ".join(ion_ion_pairs) if ion_ion_pairs  else "none"

    density_desc = {
        "global":       "bulk density of the simulation cell",
        "local":        "local z-slice density",
        "water_region": "density of the water-containing region",
    }.get(density_mode, density_mode)

    n_r = int(round(r_max / dr))

    pairing_str = ""
    if ion_pair_radii:
        pairs_listed = "; ".join(
            f"{pair}: " + "/".join(_fmt(r) for r in radii) + "\u2009\u00c5"
            for pair, radii in ion_pair_radii.items()
        )
        pairing_str = (
            " Boundaries between CIP, SIP, and DSIP states were drawn at: "
            + pairs_listed + "."
        )

    blocks = [
        (
            "p",
            f"Radial distribution functions (RDFs) were calculated as a function of "
            f"z-distance from the clay surface. The simulation box was divided into "
            f"z-slabs of thickness \u0394z\u2009=\u2009{_fmt(z_bin_width)}\u2009\u00c5, "
            f"and the trajectory was sampled every {_fmt(step)} frame(s). "
            f"For each reference ion within a slab, distances to all target atoms were "
            f"computed with the minimum-image convention (PBC), retaining distances "
            f"0\u2009<\u2009r\u2009\u2264\u2009{_fmt(r_max)}\u2009\u00c5 in {n_r} "
            f"bins of width \u0394r\u2009=\u2009{_fmt(dr)}\u2009\u00c5.",
        ),
        (
            "eq",
            "(Eq.\u20091)\u2003g(r;\u2003z_k) = "
            "[\u27e8n(r,\u0394r;\u2003z_k)\u27e9 / N_ref(z_k)] / "
            "[\u03c1_ref \u00b7 4\u03c0r\u00b2\u0394r]",
        ),
        (
            "p",
            f"The pair correlation function g(r;\u2003z_k) (Eq.\u20091) was normalised "
            f"using \u03c1_ref = {density_desc} for ion\u2013water pairs, and the "
            f"local z-slice density of the target species for ion\u2013ion pairs. "
            f"Pairs analysed: ion\u2013water ({water_str}); ion\u2013ion ({ion_ion_str}).",
        ),
    ]

    if ion_ion_pairs:
        blocks.append(
            (
                "p",
                "For cation\u2013anion pairs, g(r;\u2003z) profiles were interpreted "
                "in terms of ion-pairing speciation: contact ion pairs (CIP) represent "
                "direct cation\u2013anion contact; solvent-shared ion pairs (SIP) "
                "are separated by one water molecule; and doubly solvent-separated ion "
                "pairs (DSIP) are separated by two solvent layers."
                + pairing_str
                + " Pairing boundaries were indicated as dashed lines on the g(r;\u2003z) "
                "heat maps.",
            )
        )

    return blocks


def write_rdf_vs_z_methods(
    results_obj=None,
    params: Optional[Dict[str, Any]] = None,
    output_dir: Optional[Path] = None,
    show_in_notebook: bool = True,
    save_docx: bool = True,
):
    """Convenience wrapper: write RDF vs z methods for notebook use.

    Parameters
    ----------
    results_obj : ignored
    params : dict
        Keys drawn from ``calculate_rdf_vs_z`` and ``plot_rdf_vs_z``.
        Key fields: ion_pairs, r_max, dr, z_bin_width, step, density_mode,
        ion_pair_radii (dict {pair: [r1, r2, r3]}).
    output_dir : Path-like, optional
    show_in_notebook : bool
    save_docx : bool

    Returns
    -------
    (text, saved_path)
    """
    del results_obj
    if params is None:
        params = {}

    md = build_rdf_vs_z_methods_text(params)

    if show_in_notebook:
        blocks = build_rdf_vs_z_methods_manuscript_blocks(params)
        _show_methods_as_white_page(
            title="Methods: Radial Distribution Functions vs. z-Position",
            blocks=blocks,
        )

    saved_path = None
    if save_docx:
        out_dir = Path(output_dir) if output_dir else Path.cwd()
        out_dir.mkdir(parents=True, exist_ok=True)

        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        saved_path = out_dir / f"rdf_vs_z_methods_{ts}.docx"

        doc = Document()

        for heading_style_name in ("Heading 1", "Heading 2"):
            if heading_style_name in doc.styles:
                doc.styles[heading_style_name].font.name = "Times New Roman"
                doc.styles[heading_style_name].font.color.rgb = RGBColor(0, 0, 0)

        heading = doc.add_heading(
            "Methods: Radial Distribution Functions vs. z-Position", level=1
        )
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in heading.runs:
            run.font.name = "Times New Roman"
            run.font.color.rgb = RGBColor(0, 0, 0)

        for kind, text in build_rdf_vs_z_methods_manuscript_blocks(params):
            line = _sanitize_xml_text(text)
            if not line:
                continue
            p = doc.add_paragraph(line)
            if kind == "eq":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.space_after = Pt(8)
            for run in p.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)

        doc.styles["Normal"].font.name = "Times New Roman"
        doc.styles["Normal"].font.size = Pt(12)

        doc.save(str(saved_path))
        print(f"Word file saved -> {saved_path}")

    return md, saved_path


def write_all_methods(
    params: Optional[Dict[str, Any]] = None,
    output_dir: Optional[Path] = None,
    show_in_notebook: bool = True,
    save_docx: bool = True,
    doc_title: str = "Methods: Solvation-Shell Analysis",
):
    """Write a combined methods summary for all analyses performed in the notebook.

    Each active sub-dict is rendered as its own titled section inside one
    HTML white-page (notebook) and one combined Word document.

    Parameters
    ----------
    params : dict with optional sub-dicts:
        - ``"water_spatial"``   → passed to water spatial projection builders
        - ``"water_dipole_z"``  → passed to water dipole z-profile builders
        - ``"water_dipole_xy"`` → passed to water dipole XY map builders
        - ``"ion_spatial_xy"``  → passed to ion spatial XY distribution builders
        - ``"electrostatic"``   → passed to electrostatic potential builders
        - ``"coordination"``    → passed to coordination-shells builders
        - ``"rdf_vs_z"``        → passed to RDF vs z builders
    output_dir : Path-like, optional.  Defaults to cwd.
    show_in_notebook : bool
    save_docx : bool
    doc_title : str
        Title printed at the top of the combined document / HTML page.

    Returns
    -------
    (combined_markdown : str, saved_path : Path | None)
    """
    if params is None:
        params = {}

    SECTIONS = [
        (
            "water_spatial",
            "Water Spatial Distribution \u2014 Planar XY Maps",
            build_water_spatial_methods_manuscript_blocks,
            build_water_spatial_methods_text,
        ),
        (
            "water_dipole_z",
            "Water Dipole Orientation \u2014 z-Profile",
            build_water_dipole_z_methods_manuscript_blocks,
            build_water_dipole_z_methods_text,
        ),
        (
            "water_dipole_xy",
            "Water Dipole Orientation \u2014 Planar XY Maps",
            build_water_dipole_xy_methods_manuscript_blocks,
            build_water_dipole_xy_methods_text,
        ),
        (
            "ion_spatial_xy",
            "Ion Spatial Distribution \u2014 Planar XY Maps",
            build_ion_spatial_xy_methods_manuscript_blocks,
            build_ion_spatial_xy_methods_text,
        ),
        (
            "electrostatic",
            "Electrostatic Potential Spatial Distribution",
            build_electrostatic_potential_methods_manuscript_blocks,
            build_electrostatic_potential_methods_text,
        ),
        (
            "coordination",
            "Ion Hydration Shells \u2014 Spatial Analysis and 3D Surfaces",
            build_coordination_shells_methods_manuscript_blocks,
            build_coordination_shells_methods_text,
        ),
        (
            "rdf_vs_z",
            "Radial Distribution Functions vs. z-Position",
            build_rdf_vs_z_methods_manuscript_blocks,
            build_rdf_vs_z_methods_text,
        ),
    ]

    combined_blocks: list = []
    md_parts: list = []

    for key, section_title, blocks_fn, text_fn in SECTIONS:
        sec_params = params.get(key)
        if not sec_params:
            continue
        combined_blocks.append(("h2", section_title))
        combined_blocks.extend(blocks_fn(sec_params))
        md_parts.append(text_fn(sec_params))

    combined_md = "\n\n---\n\n".join(md_parts)

    if show_in_notebook and combined_blocks:
        _show_methods_as_white_page(title=doc_title, blocks=combined_blocks)

    saved_path = None
    if save_docx and combined_blocks:
        out_dir = Path(output_dir) if output_dir else Path.cwd()
        out_dir.mkdir(parents=True, exist_ok=True)

        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        saved_path = out_dir / f"all_methods_{ts}.docx"

        doc = Document()

        for heading_style_name in ("Heading 1", "Heading 2"):
            if heading_style_name in doc.styles:
                doc.styles[heading_style_name].font.name = "Times New Roman"
                doc.styles[heading_style_name].font.color.rgb = RGBColor(0, 0, 0)

        h1 = doc.add_heading(doc_title, level=1)
        h1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in h1.runs:
            run.font.name = "Times New Roman"
            run.font.color.rgb = RGBColor(0, 0, 0)

        for kind, text in combined_blocks:
            line = _sanitize_xml_text(str(text))
            if not line:
                continue
            if kind == "h2":
                h2 = doc.add_heading(line, level=2)
                h2.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in h2.runs:
                    run.font.name = "Times New Roman"
                    run.font.color.rgb = RGBColor(0, 0, 0)
            elif kind == "eq":
                p = doc.add_paragraph(line)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)
            else:
                p = doc.add_paragraph(line)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.space_after = Pt(8)
                for run in p.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(11)

        doc.styles["Normal"].font.name = "Times New Roman"
        doc.styles["Normal"].font.size = Pt(12)

        doc.save(str(saved_path))
        print(f"Combined methods Word file saved -> {saved_path}")

    return combined_md, saved_path


__all__ = [
    "MethodsWriter",
    "build_water_spatial_methods_text",
    "build_water_dipole_z_methods_text",
    "build_water_dipole_z_methods_manuscript_blocks",
    "build_water_dipole_xy_methods_text",
    "build_water_dipole_xy_methods_manuscript_blocks",
    "build_ion_spatial_xy_methods_text",
    "build_ion_spatial_xy_methods_manuscript_blocks",
    "build_electrostatic_potential_methods_text",
    "build_electrostatic_potential_methods_manuscript_blocks",
    "build_coordination_shells_methods_text",
    "build_coordination_shells_methods_manuscript_blocks",
    "build_rdf_vs_z_methods_text",
    "build_rdf_vs_z_methods_manuscript_blocks",
    "get_default_methods_writer",
    "write_water_spatial_methods",
    "write_water_dipole_methods",
    "write_ion_spatial_xy_methods",
    "write_electrostatic_potential_methods",
    "write_coordination_shells_methods",
    "write_rdf_vs_z_methods",
    "write_all_methods",
]
